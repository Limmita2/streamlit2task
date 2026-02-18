# -*- coding: utf-8 -*-
"""
Додаток для створення досьє на особу з PDF файлів ДМС
Підтримує завантаження декількох файлів (PDF і текстові)
"""

import streamlit as st
import os
import fitz  # PyMuPDF
import base64
from io import BytesIO
from PIL import Image
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

# Сторінка на всю ширину
st.set_page_config(page_title="Person PDF Matcher", page_icon="👥", layout="wide")

st.title("👥 Створення досьє на особу з PDF")
st.markdown("""
Завантажте PDF або текстові файли для витягування інформації про особу та створення досьє.
""")

# Ініціалізація session_state
if 'person_files_data' not in st.session_state:
    st.session_state['person_files_data'] = []
if 'person_manual_entries' not in st.session_state:
    st.session_state['person_manual_entries'] = []


def extract_dms_info_from_pdf(pdf_bytes):
    """Витягує інформацію з PDF файлу ДМС"""
    try:
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        srt_date = ''
        for page_num in range(doc.page_count):
            page = doc.load_page(page_num)
            srt_date += page.get_text()

        x = re.split("\n", srt_date)

        el = 'ІНФОРМАЦІЯ ПРО ОСОБУ'
        if el not in x:
            doc.close()
            return None, f"Файл не містить очікувану інформацію"

        # Збір інформації про особу
        obj_people = {
            'fio': '',
            'data': '',
            'birthplace': '',
            'tel': 'невідомо',
            'adress': 'невідомо',
            'uhzp': 'невідомо',
            'iphp': 'невідомо',
            'documents': []
        }

        try:
            obj_inedx = x.index('Прізвище')
            obj_people['fio'] = f'{x[obj_inedx + 1]} {x[obj_inedx + 3]} {x[obj_inedx + 5]}'
        except (ValueError, IndexError):
            pass

        try:
            str_data = x[obj_inedx + 6].split(' ')
            obj_people['data'] = str_data[2] if len(str_data) > 2 else ''
        except (ValueError, IndexError):
            pass

        try:
            obj_inedx = x.index('Телефон')
            obj_people['tel'] = x[obj_inedx + 1]
        except ValueError:
            pass

        try:
            obj_inedx = x.index('УНЗР')
            obj_people['uhzp'] = x[obj_inedx + 1]
        except ValueError:
            pass

        try:
            obj_inedx = x.index('РНОКПП')
            obj_people['iphp'] = x[obj_inedx + 1]
        except ValueError:
            pass

        def get_address(vol, date_reper):
            try:
                index_start = vol.index(date_reper[0])
                index_end = vol.index(date_reper[1])
                addr = ''
                for n in range(index_start + 1, index_end - 1):
                    addr += vol[n] + ' '

                # Форматування адреси
                addr = addr.title()
                verification = ['М.', 'Вулиця', 'Район', 'Смт', 'Кв.', 'Буд.', 'Область', 'С.', 'Вул.', "Пров.",
                              "Проспект.", "М-Н", "С-Ще", "Площа", "Проспект."]

                for slovo in addr.split():
                    if re.search(r'\d{5}', slovo) is not None:
                        addr = addr.replace(slovo, '')

                for ver in verification:
                    addr = addr.replace(ver, ver.lower())

                return addr.replace('/', ', ').strip()
            except (ValueError, IndexError):
                return 'невідомо'

        obj_people['adress'] = get_address(x, ['перебування', 'Номер'])
        obj_people['birthplace'] = get_address(x, ['Місце народження', 'перебування'])

        # Обробка документів
        teloArr = ['Паспорт громадянина України',
                  "Паспорт(и) громадянина України для виїзду за кордон",
                  'Свідоцтво про народження']

        def find_docs(arrX, index, doc_type_idx, exclude_indices):
            for w in range(index, len(arrX)):
                is_other = False
                for ex_idx in exclude_indices:
                    if arrX[w] == teloArr[ex_idx]:
                        is_other = True
                        break
                if is_other:
                    break

                if arrX[w] == "Номер":
                    if w + 4 < len(arrX) and arrX[w + 3] == "Дійсний до:":
                        obj_people['documents'].append(f"{teloArr[doc_type_idx]} {arrX[w + 1]} дійсний до: {arrX[w + 4]}")
                    elif w + 5 < len(arrX) and arrX[w + 5] == "Дійсний до:":
                        obj_people['documents'].append(f"{teloArr[doc_type_idx]} {arrX[w + 1]} від {arrX[w + 3]} дійсний до: {arrX[w + 5]}")

        for i in range(len(teloArr)):
            if teloArr[i] in x:
                idx = x.index(teloArr[i])
                others = [j for j in range(len(teloArr)) if j != i]
                find_docs(x, idx, i, others)

        # Витягування фото
        photo_bytes = None
        for img in doc.get_page_images(0):
            xref = img[0]
            pix = fitz.Pixmap(doc, xref)
            if pix.colorspace.n > 4:
                pix = fitz.Pixmap(fitz.csRGB, pix)
            photo_bytes = pix.tobytes("png")
            break

        doc.close()

        return obj_people, photo_bytes, None

    except Exception as e:
        return None, None, f"Помилка обробки PDF: {str(e)}"


def create_dossier_docx(person_data):
    """Створює DOCX документ досьє з даними про особу"""
    doc = Document()

    # Налаштування полів сторінки
    section = doc.sections[0]
    section.top_margin = Inches(2)
    section.bottom_margin = Inches(2)
    section.left_margin = Inches(3)
    section.right_margin = Inches(1.5)

    # Стиль
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(14)

    # Заголовок на блакитному фоні
    t = doc.add_table(rows=1, cols=1)
    t.width = Inches(6.5)
    cell = t.rows[0].cells[0]

    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), '9BC2E6')
    cell._element.get_or_add_tcPr().append(shading_elm)

    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run("       ІНФОРМАЦІЯ З ДМС")
    run.bold = True
    run.italic = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0, 0, 0)

    # Таблиця для фото та даних
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(0)

    table = doc.add_table(rows=1, cols=2)
    table.autofit = False

    left_cell = table.rows[0].cells[0]
    left_cell.width = Inches(2.0)

    right_cell = table.rows[0].cells[1]
    right_cell.width = Inches(4.5)
    right_cell.vertical_alignment = 1

    # Фото зліва
    photo_bytes = person_data.get('photo_bytes')
    if photo_bytes:
        p = left_cell.paragraphs[0]
        run = p.add_run()
        run.add_picture(BytesIO(photo_bytes), width=Inches(1.8))

    # Дані справа
    for para in right_cell.paragraphs:
        para.clear()

    p = right_cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # ПІБ
    if person_data.get('fio'):
        r = p.add_run(f"{person_data['fio']}\n")
        r.font.name = 'Times New Roman'
        r.font.size = Pt(14)
        r.bold = True

    # Дата народження
    if person_data.get('data'):
        r = p.add_run(f"Дата народження: {person_data['data']}\n")
        r.font.name = 'Times New Roman'
        r.font.size = Pt(14)

    # Місце народження
    if person_data.get('birthplace'):
        r = p.add_run(f"Місце народження: {person_data['birthplace']}\n")
        r.font.name = 'Times New Roman'
        r.font.size = Pt(14)

    # Телефон
    if person_data.get('tel') and person_data.get('tel') != 'невідомо':
        r = p.add_run(f"Телефон: {person_data['tel']}\n")
        r.font.name = 'Times New Roman'
        r.font.size = Pt(14)

    # УНЗР
    if person_data.get('uhzp') and person_data.get('uhzp') != 'невідомо':
        r = p.add_run(f"УНЗР: {person_data['uhzp']}\n")
        r.font.name = 'Times New Roman'
        r.font.size = Pt(14)

    # РНОКПП
    if person_data.get('iphp') and person_data.get('iphp') != 'невідомо':
        r = p.add_run(f"РНОКПП: {person_data['iphp']}\n")
        r.font.name = 'Times New Roman'
        r.font.size = Pt(14)

    # Адреса
    if person_data.get('adress') and person_data.get('adress') != 'невідомо':
        r = p.add_run(f"Адреса: {person_data['adress']}\n")
        r.font.name = 'Times New Roman'
        r.font.size = Pt(14)
        r.font.color.rgb = RGBColor(56, 86, 35)

    # Документи
    if person_data.get('documents'):
        p = doc.add_paragraph()
        r = p.add_run("ДОКУМЕНТИ:\n")
        r.bold = True
        r.font.name = 'Times New Roman'
        r.font.size = Pt(14)

        for doc_str in person_data['documents']:
            r = p.add_run(f"• {doc_str}\n")
            r.font.name = 'Times New Roman'
            r.font.size = Pt(14)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


# --- ОСНОВНИЙ ІНТЕРФЕЙС ---

col_left, col_right = st.columns([1, 1])

with col_left:
    st.subheader("📁 Завантаження файлів")

    # Завантаження файлів (PDF і текстові)
    uploaded_files = st.file_uploader(
        "Завантажте файли (PDF або текстові)",
        type=['pdf', 'txt'],
        accept_multiple_files=True,
        help="Можна завантажити кілька файлів одночасно"
    )

    if uploaded_files:
        st.write(f"🔍 Вибрано файлів: **{len(uploaded_files)}**")
        for f in uploaded_files:
            st.write(f"   • `{f.name}` ({f.type})")

    # Обробка завантажених файлів
    if uploaded_files:
        if st.button("🔄 Обробити файли", type="primary", key="process_person_files"):
            with st.spinner("Обробка файлів..."):
                all_person_data = []

                for uploaded_file in uploaded_files:
                    try:
                        st.write(f"⏳ Обробка `{uploaded_file.name}`...")

                        # Визначаємо тип файлу
                        file_ext = os.path.splitext(uploaded_file.name)[1].lower()

                        if file_ext == '.txt':
                            # Текстовий файл
                            content = uploaded_file.read().decode('utf-8')
                            # Простий парсинг для текстового формату
                            person_info = {}
                            lines = content.split('\n')
                            for line in lines:
                                if ':' in line:
                                    parts = line.split(':', 1)
                                    if len(parts) == 2:
                                        key = parts[0].strip()
                                        value = parts[1].strip()
                                        person_info[key] = value

                            if person_info:
                                all_person_data.append({
                                    'fio': person_info.get('Прізвище', ''),
                                    'data': person_info.get('Дата народження', ''),
                                    'birthplace': person_info.get('Місце народження', ''),
                                    'tel': person_info.get('Телефон', 'невідомо'),
                                    'adress': person_info.get('Адреса', 'невідомо'),
                                    'uhzp': person_info.get('УНЗР', 'невідомо'),
                                    'iphp': person_info.get('РНОКПП', 'невідомо'),
                                    'documents': [],
                                    'source': 'file',
                                    'filename': uploaded_file.name
                                })
                                st.success(f"✅ `{uploaded_file.name}` - текстовий файл оброблено")

                        elif file_ext == '.pdf':
                            # PDF файл - використовуємо PyMuPDF
                            pdf_bytes = uploaded_file.read()
                            dms_info, photo_bytes, error = extract_dms_info_from_pdf(pdf_bytes)

                            if error:
                                st.error(f"❌ `{uploaded_file.name}`: {error}")
                            elif dms_info:
                                dms_info['source'] = 'file'
                                dms_info['filename'] = uploaded_file.name
                                all_person_data.append(dms_info)
                                st.success(f"✅ `{uploaded_file.name}` - PDF оброблено")
                            else:
                                st.warning(f"⚠️ `{uploaded_file.name}` - не вдалося витягнути дані")

                        else:
                            st.warning(f"⚠️ Невідомий формат файлу `{uploaded_file.name}`")

                    except Exception as e:
                        st.error(f"❌ Помилка обробки `{uploaded_file.name}`: {str(e)}")

                if all_person_data:
                    st.session_state['person_files_data'] = all_person_data
                    st.success(f"✅ Всього оброблено: {len(all_person_data)} записів")
                    st.rerun()
                elif not all_person_data:
                    st.warning("⚠️ Не вдалося витягнути дані з жодного файлу. Перевірте формат.")

    st.markdown("---")
    st.markdown("##### **Або додати вручну:**")

    # Кнопка додавання запису
    if st.button("➕ Додати запис (ручний ввід)", key="add_manual_person"):
        st.session_state['person_manual_entries'].append({
            'text': '',
            'source': 'manual'
        })
        st.rerun()

    # Показуємо вручну додані записи
    if st.session_state.get('person_manual_entries'):
        st.markdown("**Ручний ввід:**")

        for idx in range(len(st.session_state['person_manual_entries'])):
            item = st.session_state['person_manual_entries'][idx]

            col1, col2 = st.columns([2, 1])

            with col1:
                # Текстове поле
                text_key = f"manual_person_text_{idx}"
                new_text = st.text_area(
                    f"Запис #{idx + 1}:",
                    value=item.get('text', ''),
                    key=text_key,
                    height=200
                )
                st.session_state['person_manual_entries'][idx]['text'] = new_text

            with col2:
                # Кнопка видалення
                if st.button(f"❌ Видалити #{idx + 1}", key=f"delete_manual_person_{idx}"):
                    st.session_state['person_manual_entries'].pop(idx)
                    st.rerun()

with col_right:
    st.subheader("📋 Результати обробки")

    # Об'єднуємо дані з файлів та ручного вводу
    all_results = []

    # Додаємо дані з файлів
    for item in st.session_state.get('person_files_data', []):
        all_results.append(item)

    # Додаємо дані з ручного вводу
    for item in st.session_state.get('person_manual_entries', []):
        if item.get('text'):
            # Простий парсинг для ручного введення
            person_info = {}
            lines = item['text'].split('\n')
            for line in lines:
                if ':' in line:
                    parts = line.split(':', 1)
                    if len(parts) == 2:
                        key = parts[0].strip()
                        value = parts[1].strip()
                        person_info[key] = value

            if person_info:
                all_results.append(person_info)

    if not all_results:
        st.info("👆 Завантажте файли або додайте дані вручну")
    else:
        st.write(f"**Знайдено записів: {len(all_results)}**")

        # Відображення результатів
        for idx, item in enumerate(all_results):
            with st.expander(f"👤 Особа #{idx + 1}", expanded=False):
                col1, col2 = st.columns([2, 1])

                with col1:
                    st.write("**Поля:**")
                    for key, value in item.items():
                        if key not in ['source', 'filename', 'photo_bytes'] and value:
                            st.write(f"• **{key}:** {value}")

                    if item.get('source') == 'file':
                        st.write(f"• **Джерело:** Файл `{item.get('filename', '')}`")

                with col2:
                    st.write("**Дії:**")

                    # Кнопка видалення
                    if 'filename' in item:
                        delete_key = f"delete_person_{idx}_{item.get('filename', '')}"
                    else:
                        delete_key = f"delete_person_manual_{idx}"

                    if st.button(f"🗑️ Видалити", key=delete_key):
                        if 'filename' in item:
                            # Видаляємо з файлів
                            for i, f_item in enumerate(st.session_state.get('person_files_data', [])):
                                if f_item.get('filename') == item.get('filename'):
                                    st.session_state['person_files_data'].pop(i)
                                    break
                        else:
                            # Видаляємо з ручного вводу
                            st.session_state['person_manual_entries'].pop(idx)
                        st.rerun()

                    # Кнопка завантаження DOCX
                    fio = item.get('fio', 'Особа')
                    filename_base = re.sub(r'[^\w\s-]', '', fio).strip() or 'Person'

                    docx_data = create_dossier_docx(item)

                    st.download_button(
                        label="💾 Завантажити DOCX",
                        data=docx_data,
                        file_name=f"{filename_base}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )


# Кнопка очищення
st.markdown("---")
if st.button("🧹 Очистити всі дані"):
    st.session_state['person_files_data'] = []
    st.session_state['person_manual_entries'] = []
    st.rerun()
