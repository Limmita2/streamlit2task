# -*- coding: utf-8 -*-
"""
Модуль для обробки даних про транспортні засоби (НАІС ТЗ)
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import io
import os
try:
    from .image_search import get_car_image
except ImportError:
    # Если относительный импорт не работает, используем абсолютный
    from image_search import get_car_image


def append_car_to_doc(doc: Document, car_data: list, header_name="АМТ (НАІС)"):
    """
    Додає секцію з даними про транспортні засоби до існуючого документа.

    Args:
        doc: Document об'єкт для додавання даних
        car_data: Список словників з даними про ТЗ
        header_name: Назва заголовка блоку
    """

    if not car_data or len(car_data) == 0:
        return

    # Створюємо горизонтальну смугу на всю ширину сторінки
    separator_table = doc.add_table(rows=1, cols=1)
    separator_table.width = Inches(6.5)
    separator_cell = separator_table.rows[0].cells[0]

    # Налаштування світло-блакитного фону (#9BC2E6)
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), '9BC2E6')
    separator_cell._element.get_or_add_tcPr().append(shading_elm)

    # Прибираємо границі
    tcPr = separator_cell._element.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for border in ['top', 'left', 'bottom', 'right']:
        b = OxmlElement(f'w:{border}')
        b.set(qn('w:val'), 'none')
        tcBorders.append(b)
    tcPr.append(tcBorders)

    # Додаємо текст на блакитній смузі
    p_separator = separator_cell.paragraphs[0]
    p_separator.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_separator.paragraph_format.space_before = Pt(0)
    p_separator.paragraph_format.space_after = Pt(0)
    run_separator = p_separator.add_run("       " + header_name)
    run_separator.bold = True
    run_separator.italic = True
    run_separator.font.size = Pt(14)
    run_separator.font.color.rgb = RGBColor(0, 0, 0)  # Чорний текст
    run_separator.font.name = 'Times New Roman'

    # Додаємо таблицю для кожного ТЗ з нумерацією
    for idx, car in enumerate(car_data):
        # Заголовок з нумерацією ТЗ
        tz_header = doc.add_paragraph()
        tz_header.paragraph_format.space_before = Pt(0)
        tz_header.paragraph_format.space_after = Pt(0)
        r = tz_header.add_run(f"ТЗ #{idx + 1}")
        r.font.name = 'Times New Roman'
        r.font.size = Pt(12)
        r.bold = True
        r.font.color.rgb = RGBColor(100, 100, 100)

        # Створюємо таблицю для фото та інформації
        # spacer = doc.add_paragraph()
        # spacer.paragraph_format.space_before = Pt(0)
        # spacer.paragraph_format.space_after = Pt(0)

        table = doc.add_table(rows=1, cols=2)
        table.autofit = False

        left_cell = table.rows[0].cells[0]
        left_cell.width = Inches(2.0)

        right_cell = table.rows[0].cells[1]
        right_cell.width = Inches(4.5)
        right_cell.vertical_alignment = 1

        # Отримуємо зображення автомобіля за кольором та маркою/моделлю
        car_image_bytes = None
        brand = car.get('марка', '')
        model = car.get('модель', '')
        color = car.get('колір', '')
        year = car.get('рік_випуску', '')

        if brand or color:
            car_image_bytes = get_car_image(brand=brand, model=model, color=color, year=year)

        # Додаємо зображення в ліву клітинку, якщо воно доступне
        if car_image_bytes:
            try:
                paragraph = left_cell.paragraphs[0]
                run = paragraph.add_run()
                run.add_picture(io.BytesIO(car_image_bytes), width=Inches(1.8))
            except Exception as e:
                print(f"Помилка додавання зображення: {e}")
                
        # Очищаємо стандартні параграфи в правій клітинці
        for para in right_cell.paragraphs:
            para.clear()

        # Додаємо інформацію в праву клітинку
        p = right_cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Номерний знак - жирним
        if car.get('номерний_знак'):
            r = p.add_run(f"Номерний знак: {car['номерний_знак']}\n")
            r.font.name = 'Times New Roman'
            r.font.size = Pt(14)
            r.bold = True

        # Марка/Модель
        if car.get('марка') or car.get('модель'):
            brand_model = ""
            if car.get('марка'):
                brand_model = car['марка']
            if car.get('модель'):
                brand_model += " " + car['модель']
            r = p.add_run(f"Марка/Модель: {brand_model.strip()}\n")
            r.font.name = 'Times New Roman'
            r.font.size = Pt(14)

        # VIN
        if car.get('vin'):
            r = p.add_run(f"VIN: {car['vin']}\n")
            r.font.name = 'Times New Roman'
            r.font.size = Pt(14)

        # Колір
        if car.get('колір'):
            r = p.add_run(f"Колір: {car['колір']}\n")
            r.font.name = 'Times New Roman'
            r.font.size = Pt(14)

        # Рік випуску
        if car.get('рік_випуску'):
            r = p.add_run(f"Рік випуску: {car['рік_випуску']}\n")
            r.font.name = 'Times New Roman'
            r.font.size = Pt(14)

        # Місце реєстрації
        if car.get('місце_реєстрації'):
            r = p.add_run(f"Місце реєстрації: {car['місце_реєстрації']}")
            r.font.name = 'Times New Roman'
            r.font.size = Pt(14)
            r.font.color.rgb = RGBColor(56, 86, 35)

        # # Додаємо порожній рядок між ТЗ (крім останнього)
        # if idx < len(car_data) - 1:
        #     doc.add_paragraph().paragraph_format.space_after = Pt(0)
