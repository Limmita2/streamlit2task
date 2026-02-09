# -*- coding: utf-8 -*-
"""
Модуль для обробки даних про нерухомість з PDF файлів Реєстру Нерухомості
"""

import pdfplumber
import re
import warnings
import logging
from io import BytesIO
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

logging.getLogger("pdfminer").setLevel(logging.ERROR)
warnings.filterwarnings("ignore", category=UserWarning)
warnings.filterwarnings("ignore", category=FutureWarning)


def clean_text(text):
    if not text:
        return ""
    return re.sub(r'\s+', ' ', text).strip()


def normalize_apostrophes(text):
    if not text:
        return text
    replacements = [
        ('\u2019', "'"),
        ('\u2018', "'"),
        ('\u201B', "'"),
        ('\u02BC', "'"),
    ]
    for old, new in replacements:
        text = text.replace(old, new)
    return text


def parse_real_estate_pdf(uploaded_file):
    try:
        full_text = ""
        with pdfplumber.open(uploaded_file) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    full_text += text + "\n"

        if not full_text or len(full_text.strip()) < 50:
            return None, "Не вдалося прочитати текст з файлу."

        full_text = normalize_apostrophes(full_text)
        results = []
        lines = full_text.split('\n')

        i = 0
        while i < len(lines):
            line = normalize_apostrophes(lines[i].strip())

            # --- Секція обтяжень ---
            if "актуальна інформація про державну реєстрацію обтяжень" in line.lower():
                enc_data = {}
                i += 1

                while i < len(lines):
                    line = normalize_apostrophes(lines[i].strip())

                    # Проверяем, не начался ли новый раздел (любой другой)
                    if "актуальна інформація про" in line.lower() and "державну реєстрацію обтяжень" not in line.lower():
                        break

                    if line.lower().startswith("підстава внесення запису:"):
                        basis_parts = []
                        if ':' in line:
                            basis_parts.append(clean_text(line.split(':', 1)[1]))
                        i += 1
                        while i < len(lines):
                            next_line = normalize_apostrophes(lines[i].strip())
                            # Останавливаемся только при начале нового типа информации
                            if "вид обтяження:" in next_line.lower() or ("актуальна інформація про" in next_line.lower() and "державну реєстрацію обтяжень" not in next_line.lower()):
                                break
                            basis_parts.append(next_line)
                            i += 1
                        enc_data["Підстава внесення запису"] = ' '.join(basis_parts).strip()
                        continue

                    elif line.lower().startswith("вид обтяження:"):
                        enc_data["Вид обтяження"] = clean_text(line.split(':', 1)[1]) if ':' in line else ""
                        i += 1
                        continue

                    i += 1

                if enc_data:
                    results.append(enc_data)

            # --- Секція об'єкта ---
            elif "актуальна інформація про об'єкт речових прав" in line.lower():
                # Створюємо новий об'єкт нерухомості
                current_obj = {}
                registration_dates = []  # Зберігаємо всі дати реєстрації
                shares = []  # Зберігаємо всі частки
                
                i += 1
                while i < len(lines):
                    line = normalize_apostrophes(lines[i].strip())

                    # Якщо знаходимо новий об'єкт - виходимо
                    if "актуальна інформація про об'єкт речових прав" in line.lower():
                        break

                    # Якщо знаходимо нову інформацію про речове право - додаємо дату реєстрації
                    if "актуальна інформація про речове право" in line.lower():
                        i += 1
                        continue

                    # Тип об'єкта
                    if line.lower().replace(':', '') == "тип об'єкта" or line.lower().startswith("тип об'єкта:") or line.lower().startswith("тип обєкта:"):
                        if "Тип об'єкта" not in current_obj:
                            saved_i = i
                            value = ""
                            if line.lower().replace(':', '') == "тип об'єкта" or line.lower().replace(':', '') == "тип обєкта":
                                if i + 1 < len(lines):
                                    next_line = normalize_apostrophes(lines[i + 1].strip())
                                    if ':' not in next_line or next_line.lower().startswith('так'):
                                        value = next_line
                                        i += 1
                            elif ':' in line:
                                value = clean_text(line.split(':', 1)[1])
                            if value:
                                # Clean up the type to remove extra information
                                value = value.replace('житлової нерухомості', '').strip()
                                if value.endswith(','):
                                    value = value[:-1].strip()
                                # Split by commas to get just the main type
                                if ',' in value:
                                    value = value.split(',')[0].strip()
                                current_obj["Тип об'єкта"] = value
                            else:
                                i = saved_i
                        i += 1
                        continue

                    # Кадастровий номер
                    elif line.lower().replace(':', '') == "кадастровий номер" or line.lower().startswith("кадастровий номер:"):
                        if "Кадастровий номер" not in current_obj:
                            saved_i = i
                            value = ""
                            if line.lower().replace(':', '') == "кадастровий номер":
                                if i + 1 < len(lines):
                                    value = normalize_apostrophes(lines[i + 1].strip())
                                    i += 1
                            elif ':' in line:
                                value = clean_text(line.split(':', 1)[1])
                            if value:
                                current_obj["Кадастровий номер"] = value
                            else:
                                i = saved_i
                        i += 1
                        continue

                    # Опис об'єкта
                    elif line.lower().replace(':', '') == "опис об'єкта" or line.lower().startswith("опис об'єкта:") or line.lower().startswith("опис обєкта:"):
                        if "Опис об'єкта" not in current_obj:
                            saved_i = i
                            desc_parts = []
                            if ':' in line:
                                desc_parts.append(clean_text(line.split(':', 1)[1]))
                            i += 1
                            # Читаємо до наступного поля
                            while i < len(lines):
                                next_line = normalize_apostrophes(lines[i].strip())
                                next_lower = next_line.lower()
                                # Known fields that indicate end of description
                                if any(next_lower.startswith(f) for f in ['адреса', 'кадастровий номер', 'розмір частки', 'дата, час', 'номер відомостей']):
                                    break
                                if "актуальна інформація про об'єкт речових прав" in next_lower:
                                    break
                                if "актуальна інформація про речове право" in next_lower:
                                    break
                                if next_lower.startswith('земельні ділянки') or next_lower.startswith('кадастровий номер'):
                                    # Це вже наступний блок - закінчуємо опис
                                    break
                                desc_parts.append(next_line)
                                i += 1
                            # Join description parts but clean up by removing redundant phrases
                            full_desc = ' '.join(desc_parts)
                            # Remove redundant "Актуальна інформація про речове право" from description
                            full_desc = full_desc.replace('Актуальна інформація про речове право', '').strip()
                            current_obj["Опис об'єкта"] = clean_text(full_desc)
                            i = saved_i
                        i += 1
                        continue

                    # Адреса
                    elif line.lower().replace(':', '') == "адреса" or line.lower().startswith("адреса:"):
                        if "Адреса" not in current_obj:
                            saved_i = i
                            addr_parts = []
                            if ':' in line:
                                addr_parts.append(clean_text(line.split(':', 1)[1]))
                            i += 1
                            # Читаємо до наступного поля
                            while i < len(lines):
                                next_line = normalize_apostrophes(lines[i].strip())
                                next_lower = next_line.lower()
                                if any(next_lower.startswith(f) for f in ['опис', 'кадастровий номер', 'розмір частки', 'дата, час', 'номер відомостей']):
                                    break
                                if "актуальна інформація про об'єкт речових прав" in next_lower:
                                    break
                                if "актуальна інформація про речове право" in next_lower:
                                    break
                                if next_lower.startswith('земельні ділянки') or next_lower.startswith('кадастровий номер'):
                                    # Це вже наступний блок - закінчуємо адресу
                                    break
                                addr_parts.append(next_line)
                                i += 1
                            current_obj["Адреса"] = ' '.join(addr_parts)
                            i = saved_i
                        i += 1
                        continue

                    # Розмір частки
                    elif line.lower().startswith("розмір частки:"):
                        value = line.split(':', 1)[1] if ':' in line else ""
                        share = clean_text(value)
                        if share and share not in shares and share != "1/1":
                            shares.append(share)
                        i += 1
                        continue

                    # Дата реєстрації
                    elif line.lower().startswith("дата, час державної реєстрації:"):
                        value = line.split(':', 1)[1] if ':' in line else ""
                        clean_date = clean_text(value)
                        if clean_date and clean_date not in registration_dates:
                            registration_dates.append(clean_date)
                        i += 1
                        continue

                    i += 1

                # Додаємо інформацію до об'єкта
                if shares:
                    if len(shares) == 1:
                        current_obj["Розмір частки"] = shares[0]
                    elif len(shares) > 1:
                        current_obj["Розмір частки"] = ", ".join(shares)
                
                # Якщо є кілька дат реєстрації, додаємо найпізнішу (останню)
                if registration_dates:
                    current_obj["Дата, час державної реєстрації"] = registration_dates[-1]

                if current_obj and ("Тип об'єкта" in current_obj or "Кадастровий номер" in current_obj or "Адреса" in current_obj or "Опис об'єкта" in current_obj):
                    results.append(current_obj)
                
                # Повертаємося до попереднього рядка, щоб перевірити, чи не починався новий об'єкт
                if i < len(lines) and "актуальна інформація про об'єкт речових прав" in normalize_apostrophes(lines[i].strip()).lower():
                    i -= 1  # Повертаємося, щоб наступна ітерація почала з нового об'єкта

            i += 1

        if not results:
            return None, "Немає зареєстрованої нерухомості"

        return results, None

    except Exception as e:
        return None, f"Помилка обробки файлу: {str(e)}"


def append_real_estate_to_doc(doc: Document, real_estate_data: list):
    if not real_estate_data or len(real_estate_data) == 0:
        return

    separator_table = doc.add_table(rows=1, cols=1)
    separator_table.width = Inches(6.5)
    separator_cell = separator_table.rows[0].cells[0]

    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), '9BC2E6')
    separator_cell._element.get_or_add_tcPr().append(shading_elm)

    tcPr = separator_cell._element.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for border in ['top', 'left', 'bottom', 'right']:
        b = OxmlElement(f'w:{border}')
        b.set(qn('w:val'), 'none')
        tcBorders.append(b)
    tcPr.append(tcBorders)

    p_separator = separator_cell.paragraphs[0]
    p_separator.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_separator.paragraph_format.space_before = Pt(0)
    p_separator.paragraph_format.space_after = Pt(0)
    run_separator = p_separator.add_run("       " + "НЕРУХОМІСТЬ")
    run_separator.bold = True
    run_separator.italic = True
    run_separator.font.size = Pt(14)
    run_separator.font.color.rgb = RGBColor(0, 0, 0)
    run_separator.font.name = 'Times New Roman'

    # Process each real estate item
    for idx, item in enumerate(real_estate_data):
        # Add spacing between different real estate objects
        if idx > 0:
            doc.add_paragraph()  # Empty paragraph for spacing
        
        # Determine if this is an encumbrance record (obtyzhennya) or property record
        is_encumbrance = "Вид обтяження" in item
        
        if is_encumbrance:
            # For encumbrances, use specific order
            property_order = [
                "Вид обтяження",
                "Підстава внесення запису"
            ]
        else:
            # For property records, use standard order
            property_order = [
                "Тип об'єкта",
                "Кадастровий номер", 
                "Опис об'єкта",
                "Адреса", 
                "Розмір частки",
                "Дата, час державної реєстрації"
            ]
        
        for prop_name in property_order:
            if prop_name in item and item[prop_name]:
                value = str(item[prop_name]).strip()
                
                # Skip if value is just "так" (which indicates presence rather than actual value)
                if value.lower() == "так":
                    continue
                
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(2)

                run_key = p.add_run(f"{prop_name}: ")
                run_key.bold = True
                run_key.font.size = Pt(14)
                run_key.font.name = 'Times New Roman'

                run_value = p.add_run(value)
                run_value.font.size = Pt(14)
                run_value.font.name = 'Times New Roman'
