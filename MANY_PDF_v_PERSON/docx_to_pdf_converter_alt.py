import io
from docx import Document
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import inch
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_JUSTIFY
from io import BytesIO


def convert_docx_to_pdf(docx_bytes: bytes) -> bytes:
    """
    Альтернативная реализация конвертации DOCX в PDF с использованием reportlab.
    Эта функция анализирует DOCX-документ и воссоздает его структуру в PDF.

    Args:
        docx_bytes: Байты DOCX-документа

    Returns:
        bytes: Байты PDF-документа
    """
    # Регистрируем системные шрифты с поддержкой кириллицы
    import os
    try:
        # Пробуем Liberation Sans
        if os.path.exists('/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf'):
            pdfmetrics.registerFont(TTFont('LiberationSans', '/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf'))
            pdfmetrics.registerFont(TTFont('LiberationSans-Bold', '/usr/share/fonts/truetype/liberation/LiberationSans-Bold.ttf'))
            pdfmetrics.registerFont(TTFont('LiberationSans-Italic', '/usr/share/fonts/truetype/liberation/LiberationSans-Italic.ttf'))
            pdfmetrics.registerFont(TTFont('LiberationSans-BoldItalic', '/usr/share/fonts/truetype/liberation/LiberationSans-BoldItalic.ttf'))

            default_font = 'LiberationSans'
            bold_font = 'LiberationSans-Bold'
            italic_font = 'LiberationSans-Italic'
            bold_italic_font = 'LiberationSans-BoldItalic'
        # Пробуем DejaVu Sans
        elif os.path.exists('/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf'):
            pdfmetrics.registerFont(TTFont('DejaVuSans', '/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf'))
            pdfmetrics.registerFont(TTFont('DejaVuSans-Bold', '/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf'))
            pdfmetrics.registerFont(TTFont('DejaVuSans-Oblique', '/usr/share/fonts/truetype/dejavu/DejaVuSans-Oblique.ttf'))
            pdfmetrics.registerFont(TTFont('DejaVuSans-BoldOblique', '/usr/share/fonts/truetype/dejavu/DejaVuSans-BoldOblique.ttf'))

            default_font = 'DejaVuSans'
            bold_font = 'DejaVuSans-Bold'
            italic_font = 'DejaVuSans-Oblique'
            bold_italic_font = 'DejaVuSans-BoldOblique'
        else:
            raise Exception("Системные шрифты не найдены")
    except Exception as e:
        print(f"Помилка завантаження шрифту: {e}")
        default_font = 'Helvetica'
        bold_font = 'Helvetica-Bold'
        italic_font = 'Helvetica-Oblique'
        bold_italic_font = 'Helvetica-BoldOblique'

    # Загружаем DOCX документ
    docx_buffer = io.BytesIO(docx_bytes)
    doc = Document(docx_buffer)

    # Создаем PDF документ в памяти
    pdf_buffer = BytesIO()
    doc_template = SimpleDocTemplate(pdf_buffer, pagesize=A4)

    # Устанавливаем стили
    styles = getSampleStyleSheet()
    custom_styles = {}

    # Создаем кастомные стили
    custom_styles['TitleCenter'] = ParagraphStyle(
        'TitleCenter',
        parent=styles['Normal'],
        fontName=default_font,
        fontSize=14,
        alignment=TA_CENTER,
        spaceAfter=6,
        spaceBefore=0,
        leading=16.1  # 1,15 интервал
    )

    custom_styles['BlueHeader'] = ParagraphStyle(
        'BlueHeader',
        parent=styles['Normal'],
        fontName=bold_font,
        fontSize=14,
        alignment=TA_LEFT,
        textColor=colors.white,
        backColor=colors.HexColor('#9BC2E6'),
        leftIndent=28,  # 7 пробелов (~7*4pt)
        spaceAfter=0,
        spaceBefore=0,
        leading=16.1  # 1,15 интервал
    )

    custom_styles['NormalIndented'] = ParagraphStyle(
        'NormalIndented',
        parent=styles['Normal'],
        fontName=default_font,
        fontSize=14,
        alignment=TA_JUSTIFY,
        leftIndent=0,
        spaceAfter=2,
        spaceBefore=0,
        leading=16.1  # 1,15 интервал
    )

    # Список элементов для построения PDF
    elements = []

    # Проходим по параграфам документа
    for paragraph in doc.paragraphs:
        text = paragraph.text

        if text.strip():  # Если текст не пустой
            # Определяем стиль на основе форматирования
            style_key = 'NormalIndented'

            # Проверяем, есть ли жирный текст в параграфе
            has_bold = any(run.bold for run in paragraph.runs)
            is_centered = paragraph.alignment == 1  # CENTER alignment

            # Проверяем, содержит ли текст заголовок "АНКЕТНІ ДАНІ:" или другие заголовки
            if "АНКЕТНІ ДАНІ:" in text or "АНАЛІТИЧНИЙ ПРОФІЛЬ" in text or "на фізичну особу" in text:
                style_key = 'TitleCenter'
            elif any(word in text.upper() for word in ["АНКЕТНІ ДАНІ:", "ЄРДР", "АДРЕСА", "КОНТАКТИ", "ПОЧАТОК ДОКУМЕНТА"]):
                # Это заголовок с синим фоном
                style = ParagraphStyle(
                    'BlueHeaderTemp',
                    parent=styles['Normal'],
                    fontName=bold_font,
                    fontSize=14,
                    alignment=TA_LEFT,
                    textColor=colors.white,
                    backColor=colors.HexColor('#9BC2E6'),
                    leftIndent=28,  # 7 пробелов (~7*4pt)
                    spaceAfter=0,
                    spaceBefore=0,
                    leading=16.1
                )
                p = Paragraph(text, style)
                elements.append(p)
            else:
                # Создаем стиль с учетом форматирования
                if has_bold:
                    # Если есть жирный текст, создаем специальный стиль
                    style = ParagraphStyle(
                        'BoldStyle',
                        parent=styles['Normal'],
                        fontName=bold_font,
                        fontSize=14,
                        alignment=TA_JUSTIFY,
                        leftIndent=0,
                        spaceAfter=2,
                        spaceBefore=0,
                        leading=16.1
                    )
                    p = Paragraph(text, style)
                    elements.append(p)
                else:
                    # Обычный параграф
                    p = Paragraph(text, custom_styles[style_key])
                    elements.append(p)
        else:
            # Добавляем пустой параграф для сохранения форматирования
            elements.append(Spacer(1, 8))  # 8pt высота строки

    # Обрабатываем таблицы
    for table in doc.tables:
        # Преобразуем таблицу в формат reportlab
        data = []
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                cell_text = cell.text
                # Проверяем, есть ли в ячейке синий фон
                if cell._tc.pr.shd and cell._tc.pr.shd.val:
                    bg_color = cell._tc.pr.shd.val
                    # Если цвет близок к #9BC2E6, применяем стиль заголовка
                    if '#9BC2E6' in str(bg_color) or '9BC2E6' in str(bg_color):
                        # Создаем параграф с синим фоном
                        style = ParagraphStyle(
                            'BlueCell',
                            parent=styles['Normal'],
                            fontName=bold_font,
                            fontSize=14,
                            alignment=TA_LEFT,
                            textColor=colors.white,
                            backColor=colors.HexColor('#9BC2E6'),
                            leftIndent=28,  # 7 пробелов (~7*4pt)
                            spaceAfter=0,
                            spaceBefore=0,
                            leading=16.1
                        )
                        row_data.append(Paragraph(cell_text, style))
                    else:
                        row_data.append(cell_text)
                else:
                    row_data.append(cell_text)
            data.append(row_data)

        if data:
            # Создаем таблицу
            tbl = Table(data)

            # Проверяем, есть ли у таблицы синий заголовок
            if data and len(data) > 0:
                # Применяем стиль ко всей таблице
                tbl.setStyle(TableStyle([
                    ('FONTNAME', (0, 0), (-1, -1), default_font),
                    ('FONTSIZE', (0, 0), (-1, -1), 14),
                    ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                    ('LEFTPADDING', (0, 0), (-1, -1), 4),
                    ('RIGHTPADDING', (0, 0), (-1, -1), 4),
                    ('TOPPADDING', (0, 0), (-1, -1), 1),
                    ('BOTTOMPADDING', (0, 0), (-1, -1), 1),
                ]))

            elements.append(tbl)
            elements.append(Spacer(1, 6))  # Отступ после таблицы

    # Строим PDF
    doc_template.build(elements)

    # Возвращаем байты PDF
    pdf_buffer.seek(0)
    return pdf_buffer.getvalue()


def get_pdf_filename_from_docx(docx_filename: str) -> str:
    """
    Генерирует имя PDF-файла из имени DOCX-файла.
    
    Args:
        docx_filename: Имя DOCX-файла
        
    Returns:
        str: Имя PDF-файла
    """
    if docx_filename.lower().endswith('.docx'):
        return docx_filename[:-5] + '.pdf'
    else:
        return docx_filename + '.pdf'