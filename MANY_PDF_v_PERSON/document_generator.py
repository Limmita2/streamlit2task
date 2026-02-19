from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import inch
from reportlab.pdfgen import canvas
from reportlab.lib.utils import ImageReader
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from io import BytesIO
from PIL import Image
import io
import os
from datetime import datetime
from arkan_processor import append_border_crossing_to_doc
from dms_processor import append_dms_to_doc
from real_estate_processor import append_real_estate_to_doc
try:
    from .car_processor import append_car_to_doc
except ImportError:
    from car_processor import append_car_to_doc
try:
    from .pension_processor import process_pension_data
except ImportError:
    from pension_processor import process_pension_data


import re
import base64

EMPTY_DOSSIER_BLOCKS = [
    "АНКЕТНІ ДАНІ",
    "ДОКУМЕНТИ",
    "АДРЕСИ",
    "ПРИВАТНІ СТОРІНКИ В СОЦІАЛЬНИХ МЕРЕЖАХ",
    "ПІДПРИЄМНИЦЬКА (ТРУДОВА) ДІЯЛЬНІСТЬ",
    "НЕРУХОМЕ МАЙНО",
    "ТРАНСПОРТНІ ЗАСОБИ",
    "АДМІНІСТРАТИВНА ВІДПОВІДАЛЬНІСТЬ",
    "КРИМІНАЛЬНА ВІДПОВІДАЛЬНІСТЬ",
    "ОБЛІКИ ОВС",
    "ЗАРЕЄСТРОВАНА ЗБРОЯ",
    "ІНФОРМАЦІЯ З РЕЄСТРУ СУДОВИХ РІШЕНЬ",
    "ВИКОНАВЧІ ПРОВАДЖЕННЯ, РЕЄСТР БОРЖНИКІВ",
    "ДОВІРЕНОСТІ",
    "РОДИННІ ЗВ'ЯЗКИ",
    "ПЕРЕТИНИ ДЕРЖАВНОГО КОРДОНУ УКРАЇНИ"
]

BLOCK_MAPPING = {
    "Початок документа": "АНКЕТНІ ДАНІ",
    "Адреса": "АДРЕСИ",
    "Нерухомість": "НЕРУХОМЕ МАЙНО",
    "Нерухоме майно": "НЕРУХОМЕ МАЙНО",
    "АВТО (НАІС ТЗ)": "ТРАНСПОРТНІ ЗАСОБИ",
    "АВТО НАІС ТЗ": "ТРАНСПОРТНІ ЗАСОБИ",
    "БАЗА НАІС ТЗ": "ТРАНСПОРТНІ ЗАСОБИ",
    "авто (наіс тз)": "ТРАНСПОРТНІ ЗАСОБИ",
}


def get_filename_from_intro(data: dict) -> str:
    """
    Витягує перше слово з блоку 'Початок документа' для формування імені файлу.
    """
    content_list = data.get("Контент", [])

    for item in content_list:
        if item.get("header") == "Початок документа":
            content = item.get("content", "")
            # Витягуємо перше слово з контенту
            first_word = content.split()[0] if content.split() else "Dossier"
            # Прибираємо спеціальні символи з імені файлу
            import re
            first_word = re.sub(r'[^\w\s-]', '', first_word)
            return f"{first_word}.docx"

    return "Dossier.docx"


def generate_docx(data: dict, photo_bytes: bytes = None, border_crossing_data: list = None, dms_data: dict = None, family_data: list = None, real_estate_data: list = None, car_data: list = None, pension_data: dict = None) -> bytes:
    """
    Генерує документ Word з вибраних абзаців.
    """
    doc = Document()

    # Налаштування полів сторінки
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3)
    section.right_margin = Cm(1.5)

    # Налаштування стилів
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(14)


    BOLD_PATTERN = r'(Mарка\s*:|заявник\s*:|Марка\s*:|свідок\s*\(учасник\)\s*:|ухилянт\s*:|Вид\s*:|правопорушник\s*:|Номер\s*дозволу\s*:|телефони\s*:|[МM][іi][сc]ц[еe]\s*[нH][аa][рp][оo]дж[еe][нH]{2}я\s*:|Громадянство\s*:|затриманий\s*:|постраждалий\s*\(потерпілий\)\s*:|категорія\s*:|№\s+[А-ЯІЇ]{2,4}\s+\d+(?:\s+[А-ЯІЇ]{2}\s+\d+)?\s+від\s+\d{2}\.\d{2}\.\d{4}\s+\d{2}:\d{2}:\d{2}\s*,\s*орган:)'

    INTRO_PATTERN = r'(Місце\s*народження\s*:|Громадянство\s*:)'


    def add_bulleted_content(container, text, alignment=None, use_bullet_style=True, bold_matches=True, bold_content=False, pattern=BOLD_PATTERN, exclude_pattern=None, use_first_paragraph=False):
        """Разбивает текст по шаблону и создает маркированный список для ключевых слов."""
        if pattern:
            parts = re.split(pattern, text)
            current_p = None
            part_index = 0

            for part in parts:
                if not part:
                    continue

                # Проверяем, является ли часть ключевым словом
                if re.fullmatch(pattern, part):
                    # Если часть совпадает с исключаемым паттерном, не делаем её жирной
                    if exclude_pattern and re.fullmatch(exclude_pattern, part):
                        is_bold = False
                    else:
                        is_bold = bold_matches

                    # Начинаем новый абзац (маркированный или обычный)
                    style = 'List Bullet' if use_bullet_style else None
                    if use_first_paragraph and part_index == 0 and hasattr(container, 'paragraphs') and len(container.paragraphs) > 0:
                        current_p = container.paragraphs[0]
                    else:
                        current_p = container.add_paragraph(style=style)
                    current_p.paragraph_format.space_before = Pt(0)
                    current_p.paragraph_format.space_after = Pt(2)
                    if alignment is not None:
                        current_p.alignment = alignment

                    run = current_p.add_run(part)
                    run.bold = is_bold
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(14)
                    part_index += 1
                else:
                    if current_p is None:
                        # Если ключевых слов еще не было, создаем обычный абзац
                        if use_first_paragraph and hasattr(container, 'paragraphs') and len(container.paragraphs) > 0:
                            current_p = container.paragraphs[0]
                        else:
                            current_p = container.add_paragraph()
                        current_p.paragraph_format.space_before = Pt(0)
                        current_p.paragraph_format.space_after = Pt(2)
                        if alignment is not None:
                            current_p.alignment = alignment

                    run = current_p.add_run(part)
                    run.bold = bold_content
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(14)
        else:
            # Просто добавляем текст как обычный
            current_p = container.add_paragraph()
            current_p.paragraph_format.space_before = Pt(0)
            current_p.paragraph_format.space_after = Pt(2)
            if alignment is not None:
                current_p.alignment = alignment
            run = current_p.add_run(text)
            run.bold = bold_content
            run.font.name = 'Times New Roman'
            run.font.size = Pt(14)

    # Додаємо заголовки над блоком "АНАЛІТИЧНЕ ДОСЬЄ НА ОСОБУ"
    p_analitic_profile = doc.add_paragraph()
    p_analitic_profile.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_analitic_profile.paragraph_format.space_before = Pt(0)
    p_analitic_profile.paragraph_format.space_after = Pt(0)
    p_analitic_profile.paragraph_format.line_spacing = 1.15  # Устанавливаем межстрочный интервал 1,15
    run_analitic_profile = p_analitic_profile.add_run("АНАЛІТИЧНИЙ ПРОФІЛЬ")
    run_analitic_profile.bold = True
    run_analitic_profile.font.size = Pt(14)


    p_on_person = doc.add_paragraph()
    p_on_person.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_on_person.paragraph_format.space_before = Pt(0)
    p_on_person.paragraph_format.space_after = Pt(0)
    p_on_person.paragraph_format.line_spacing = 1.15  # Устанавливаем межстрочный интервал 1,15
    run_on_person = p_on_person.add_run("на фізичну особу")
    run_on_person.bold = True
    run_on_person.font.size = Pt(14)


    # Пустая строка после "на фізичну особу"
    empty_line_after_person = doc.add_paragraph()
    empty_line_after_person.paragraph_format.space_before = Pt(0)
    empty_line_after_person.paragraph_format.space_after = Pt(0)
    empty_line_after_person.paragraph_format.line_spacing = 1.15


    # Шукаємо вступний текст (Початок документа)
    content_list = data.get("Контент", [])
    intro_text = ""
    filtered_content = []

    for item in content_list:
        header = item.get("header", "").strip()
        if header == "Початок документа" and not intro_text:
            intro_text = item.get("content", "")
        elif dms_data and (header == "АНКЕТНІ ДАНІ:" or header == "АНКЕТНІ ДАНІ"):
            # Пропускаємо цей блок, бо ДМС його замінить
            continue
        else:
            filtered_content.append(item)

    # Пріоритет фото для всього документа: 1. ДМС, 2. Завантажене вручну, 3. Значення за замовчуванням
    final_photo_bytes = None
    if dms_data and dms_data.get('photo_bytes'):
         final_photo_bytes = dms_data['photo_bytes']
    elif photo_bytes:
         final_photo_bytes = photo_bytes
    elif os.path.exists('default_avatar.png'):
         with open('default_avatar.png', 'rb') as f:
             final_photo_bytes = f.read()

    # ЛОГІКА ЗАМІНИ АНКЕТНИХ ДАНИХ НА ДМС
    if dms_data and dms_data.get('info'):
        # Виводимо ДМС першим блоком із заголовком "ІНФОРМАЦІЯ З ДМС"
        append_dms_to_doc(doc, dms_data['info'], photo_bytes=final_photo_bytes, header_name="ІНФОРМАЦІЯ З ДМС")
    else:
        # 1. ЗАГАЛЬНИЙ ЗАГОЛОВОК (тільки якщо немає ДМС)
        t_top = doc.add_table(rows=1, cols=1)
        t_top.width = Inches(6.5)
        cell_top = t_top.rows[0].cells[0]
        
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), '9BC2E6')
        cell_top._element.get_or_add_tcPr().append(shd)
        
        p_top = cell_top.paragraphs[0]
        p_top.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p_top.paragraph_format.space_before = Pt(0)
        p_top.paragraph_format.space_after = Pt(0)
        run_top = p_top.add_run("       " + "АНКЕТНІ ДАНІ:")
        run_top.bold = True
        run_top.italic = True
        run_top.font.size = Pt(14)

        # 2. Створюємо стандартну вступну таблицю (АНКЕТНІ ДАНІ), якщо немає ДМС
        spacer = doc.add_paragraph()
        spacer.paragraph_format.space_before = Mm(3)
        spacer.paragraph_format.space_after = Mm(0)
        spacer.paragraph_format.line_spacing = 0

        table = doc.add_table(rows=1, cols=2)
        table.autofit = False
        
        left_cell = table.rows[0].cells[0]
        left_cell.width = Inches(2.0)
        if final_photo_bytes:
            paragraph = left_cell.paragraphs[0]
            run = paragraph.add_run()
            run.add_picture(BytesIO(final_photo_bytes), width=Inches(1.8))
            
        right_cell = table.rows[0].cells[1]
        right_cell.width = Inches(4.5)
        right_cell.vertical_alignment = 1

        # Очищаємо стандартні параграфи у правій клітинці перед додаванням контенту
        for para in right_cell.paragraphs:
            para.clear()

        if intro_text:
            intro_text = intro_text.replace("д.н.", "").replace("  ", " ")
            add_bulleted_content(right_cell, intro_text, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                                 use_bullet_style=False, bold_matches=True, bold_content=True, pattern=BOLD_PATTERN, exclude_pattern=INTRO_PATTERN, use_first_paragraph=True)
        else:
            title_paragraph = right_cell.paragraphs[0]
            title_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            title_run = title_paragraph.add_run("Особисте досьє")
            title_run.font.size = Pt(14)
            title_run.font.bold = True
            title_run.font.color.rgb = RGBColor(0, 0, 0)

    # Додаємо секцію нерухомості, якщо вона є (має бути другою за логікою)
    if real_estate_data:
        append_real_estate_to_doc(doc, real_estate_data)

    # Добавляем контент (вже відфільтрований без вступу)
    for item in filtered_content:
        header = item.get("header", "").strip()
        content = item.get("content", "").strip()


        if header:
            if header == "Початок документа":
                # Виводимо тільки контент як звичайний текст на початку
                if content:
                    add_bulleted_content(doc, content, pattern=None)
                    # Добавляем отступ после вводного блока
                    doc.add_paragraph().paragraph_format.space_after = Pt(6)
                continue


            # Створюємо таблицю для заголовка на блакитному фоні
            t = doc.add_table(rows=1, cols=1)
            t.width = Inches(6.5)
            cell = t.rows[0].cells[0]


            # Налаштування блакитного фону (#9BC2E6)
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), '9BC2E6')
            cell._element.get_or_add_tcPr().append(shading_elm)


            # Прибираємо границі
            tcPr = cell._element.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for border in ['top', 'left', 'bottom', 'right']:
                b = OxmlElement(f'w:{border}')
                b.set(qn('w:val'), 'none')
                tcBorders.append(b)
            tcPr.append(tcBorders)


            p_h = cell.paragraphs[0]
            p_h.alignment = WD_ALIGN_PARAGRAPH.LEFT
            # Додаємо 7 пробілів перед заголовком та робимо текст великими літерами
            run_h = p_h.add_run("       " + header.upper())
            run_h.bold = True
            run_h.italic = True
            run_h.font.size = Pt(14)
            p_h.paragraph_format.space_before = Pt(0)
            p_h.paragraph_format.space_after = Pt(0)


            paragraphs_list = content.split('\n')
            for i, p_text in enumerate(paragraphs_list):
                if p_text.strip():
                    # Применяем выравнивание по центру для всех блоков кроме "Початок документа"
                    pat = (r'(№\s+\d{24}\s+від\s+\d{2}\.\d{2}\.\d{4}\s*,\s*за\s*СТ\.|' + BOLD_PATTERN[1:] if header == "ЄРДР" else
                           r'(місце\s*проживання\s*:|' + BOLD_PATTERN[1:] if header == "Адреса" else BOLD_PATTERN)
                    p_c = add_bulleted_content(doc, p_text.strip(), alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, pattern=pat)


    # Знаходимо абзац з текстом "АНКЕТНІ ДАНІ:" і змінюємо формат порожніх абзаців перед і після нього
    paragraphs = doc.paragraphs
    for i, para in enumerate(paragraphs):
        if "АНКЕТНІ ДАНІ:" in para.text:
            # Перевіряємо, чи існує абзац перед цим
            if i > 0:
                prev_para = paragraphs[i-1]
                if not prev_para.text.strip():  # Якщо абзац порожній
                    # Встановлюємо висоту рядка точно 8 пт
                    prev_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
                    prev_para.paragraph_format.line_spacing = Pt(8)
                    # Встановлюємо розмір шрифта 8 для всіх runs у цьому абзаці
                    for run in prev_para.runs:
                        run.font.size = Pt(8)

            # Перевіряємо, чи існує абзац після цього
            if i < len(paragraphs) - 1:
                next_para = paragraphs[i+1]
                if not next_para.text.strip():  # Якщо абзац порожній
                    # Встановлюємо висоту рядка точно 8 пт
                    next_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
                    next_para.paragraph_format.line_spacing = Pt(8)
                    # Встановлюємо розмір шрифта 8 для всіх runs у цьому абзаці
                    for run in next_para.runs:
                        run.font.size = Pt(8)
            break


    # Додаємо верхній і нижній колонтитул
    section = doc.sections[0]


    # Налаштовуємо верхній колонтитул
    header = section.header
    header.is_linked_to_previous = False  # Відключаємо зв'язок з попередніми секціями


    # Додаємо перший параграф у верхній колонтитул (текст "УКА ГУНП")
    header_para1 = header.paragraphs[0]
    header_para1.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    header_run1 = header_para1.add_run("УКА ГУНП")
    header_run1.font.name = "Times New Roman"
    header_run1.font.size = Pt(12)
    header_run1.font.color.rgb = RGBColor(255, 0, 0)  # Червоний колір


    # Додаємо другий параграф у верхній колонтитул (сьогоднішня дата)
    header_para2 = header.add_paragraph()
    header_para2.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    today_date = datetime.today().strftime("%d.%m.%Y")
    header_run2 = header_para2.add_run(today_date)
    header_run2.font.name = "Times New Roman"
    header_run2.font.size = Pt(12)
    header_run2.font.color.rgb = RGBColor(255, 0, 0)  # Червоний колір


    # Налаштовуємо нижній колонтитул
    footer = section.footer
    footer.is_linked_to_previous = False  # Відключаємо зв'язок з попередніми секціями


    # Додаємо параграф у нижній колонтитул
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    footer_run = footer_para.add_run("Управління кримінального аналізу")
    footer_run.font.name = "Times New Roman"
    footer_run.font.size = Pt(12)
    footer_run.font.bold = True


    # Додаємо родинні зв'язки (ДМС родичів), якщо вони є
    if family_data:
        for member in family_data:
            # member - це словник {'relative_type': 'дружина', 'info': dms_info, 'photo_bytes': bytes}
            # або {'relative_type': 'дружина', 'manual_text': '...', 'photo_bytes': bytes}
            header = member.get('relative_type', 'РОДИЧ').upper()
            if member.get('info'):
                append_dms_to_doc(doc, member['info'], photo_bytes=member.get('photo_bytes'), header_name=f"{header} (ДМС)")
            elif member.get('manual_text'):
                # Для вручну введених даних створюємо таблицю з фото та текстом
                # 1. Заголовок на блакитному фоні
                separator_table = doc.add_table(rows=1, cols=1)
                separator_table.width = Inches(6.5)
                separator_cell = separator_table.rows[0].cells[0]

                shading_elm = OxmlElement('w:shd')
                shading_elm.set(qn('w:fill'), '9BC2E6')
                separator_cell._element.get_or_add_tcPr().append(shading_elm)

                tcPr = separator_cell._element.get_or_add_tcPr()
                tcBorders = OxmlElement('w:tcBorders')
                for border in ['top', 'left', 'bottom', 'right']:
                    b = OxmlElement(f'w:{border}')
                    b.set(qn('w:val'), 'none')
                    tcBorders.append(b)
                tcPr.append(tcBorders)

                p_separator = separator_cell.paragraphs[0]
                p_separator.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p_separator.paragraph_format.space_before = Pt(0)
                p_separator.paragraph_format.space_after = Pt(0)
                run_separator = p_separator.add_run("       " + header)
                run_separator.bold = True
                run_separator.italic = True
                run_separator.font.size = Pt(14)
                run_separator.font.color.rgb = RGBColor(0, 0, 0)
                run_separator.font.name = 'Times New Roman'

                # 2. Таблиця з фото (зліва) та текстом (справа)
                spacer = doc.add_paragraph()
                spacer.paragraph_format.space_before = Mm(3)
                spacer.paragraph_format.space_after = Mm(0)
                spacer.paragraph_format.line_spacing = 0

                table = doc.add_table(rows=1, cols=2)
                table.autofit = False

                left_cell = table.rows[0].cells[0]
                left_cell.width = Inches(2.0)
                photo_bytes = member.get('photo_bytes')
                if photo_bytes:
                    paragraph = left_cell.paragraphs[0]
                    run = paragraph.add_run()
                    run.add_picture(BytesIO(photo_bytes), width=Inches(1.8))
                else:
                    # Пробуем найти default_avatar.png в нескольких местах
                    avatar_paths = [
                        'default_avatar.png',
                        'MANY_PDF_v_PERSON/default_avatar.png',
                        os.path.join(os.path.dirname(__file__), 'default_avatar.png')
                    ]
                    avatar_path = None
                    for path in avatar_paths:
                        if os.path.exists(path):
                            avatar_path = path
                            break

                    if avatar_path:
                        with open(avatar_path, 'rb') as f:
                            paragraph = left_cell.paragraphs[0]
                            run = paragraph.add_run()
                            run.add_picture(BytesIO(f.read()), width=Inches(1.8))

                right_cell = table.rows[0].cells[1]
                right_cell.width = Inches(4.5)
                right_cell.vertical_alignment = 1

                # Додаємо текст у праву клітинку
                manual_text = member.get('manual_text', '')
                if manual_text:
                    # Розбиваємо текст на рядки та додаємо кожен як окремий абзац
                    lines = manual_text.split('\n')
                    for line_idx, line in enumerate(lines):
                        if line.strip():
                            # Для першого рядка використовуємо існуючий параграф
                            if line_idx == 0:
                                p = right_cell.paragraphs[0]
                                p.clear()
                            else:
                                p = right_cell.add_paragraph()
                            p.paragraph_format.space_before = Pt(0)
                            p.paragraph_format.space_after = Pt(2)
                            # Тільки перший рядок робимо жирним
                            if line_idx == 0:
                                run = p.add_run(line.strip())
                                run.bold = True
                            else:
                                run = p.add_run(line.strip())
                            run.font.name = 'Times New Roman'
                            run.font.size = Pt(14)

    # Додаємо секцію про транспортні засоби, якщо вона є
    if car_data:
        append_car_to_doc(doc, car_data)

    # Додаємо секцію про Пенсійний фонд, якщо вона є
    if pension_data:
        append_pension_to_doc(doc, pension_data)

    # Додаємо секцію про перетин кордону, якщо вона є
    if border_crossing_data:
        append_border_crossing_to_doc(doc, border_crossing_data)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def append_pension_to_doc(doc, pension_data: dict, add_header: bool = True):
    """
    Додає дані з Пенсійного фонду до документу Word.
    Виводить інформацію в один рядок у розділ "ПІДПРИЄМНИЦЬКА (ТРУДОВА) ДІЯЛЬНІСТЬ".

    Args:
        doc: Document об'єкт
        pension_data: dict з ключем 'formatted_line' що містить відформатований рядок
        add_header: Чи додавати заголовок (True для generate_empty_dossier, False для generate_docx)
    """
    if not pension_data or not pension_data.get('formatted_line'):
        return

    # Додаємо заголовок тільки якщо потрібно
    if add_header:
        add_block_header(doc, "ПІДПРИЄМНИЦЬКА (ТРУДОВА) ДІЯЛЬНІСТЬ")

    # Додаємо текст одним рядком
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(pension_data['formatted_line'])
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)


def add_block_header(doc, header_name: str):
    """Додає тільки заголовок блоку на блакитному фоні без порожніх рядків."""
    t = doc.add_table(rows=1, cols=1)
    t.width = Inches(6.5)
    cell = t.rows[0].cells[0]

    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), '9BC2E6')
    cell._element.get_or_add_tcPr().append(shading_elm)

    tcPr = cell._element.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for border in ['top', 'left', 'bottom', 'right']:
        b = OxmlElement(f'w:{border}')
        b.set(qn('w:val'), 'none')
        tcBorders.append(b)
    tcPr.append(tcBorders)

    p_h = cell.paragraphs[0]
    p_h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run_h = p_h.add_run("       " + header_name.upper())
    run_h.bold = True
    run_h.italic = True
    run_h.font.size = Pt(14)
    p_h.paragraph_format.space_before = Pt(0)
    p_h.paragraph_format.space_after = Pt(0)


def add_empty_block(doc, header_name: str, photo_bytes: bytes = None):
    """Додає порожній блок із заголовком на блакитному фоні та 5 порожніх рядків."""
    t = doc.add_table(rows=1, cols=1)
    t.width = Inches(6.5)
    cell = t.rows[0].cells[0]

    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), '9BC2E6')
    cell._element.get_or_add_tcPr().append(shading_elm)

    tcPr = cell._element.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for border in ['top', 'left', 'bottom', 'right']:
        b = OxmlElement(f'w:{border}')
        b.set(qn('w:val'), 'none')
        tcBorders.append(b)
    tcPr.append(tcBorders)

    p_h = cell.paragraphs[0]
    p_h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run_h = p_h.add_run("       " + header_name.upper())
    run_h.bold = True
    run_h.italic = True
    run_h.font.size = Pt(14)
    p_h.paragraph_format.space_before = Pt(0)
    p_h.paragraph_format.space_after = Pt(0)

    for _ in range(5):
        empty_p = doc.add_paragraph()
        empty_p.paragraph_format.space_before = Pt(0)
        empty_p.paragraph_format.space_after = Pt(0)


def generate_empty_dossier(photo_bytes: bytes = None, border_crossing_data: list = None,
                           dms_data: dict = None, family_data: list = None,
                           real_estate_data: list = None, car_data: list = None,
                           pension_data: dict = None,
                           filled_blocks: dict = None) -> bytes:
    """
    Генерує порожнє досьє з усіма блоками.
    
    Args:
        photo_bytes: Фото для анкетних даних
        border_crossing_data: Дані про перетин кордону
        dms_data: Дані ДМС
        family_data: Дані про родинні зв'язки
        real_estate_data: Дані про нерухомість
        car_data: Дані про транспортні засоби
        filled_blocks: Словник {header_name: content} для заповнених блоків з PDF
    
    Returns:
        bytes: DOCX файл
    """
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3)
    section.right_margin = Cm(1.5)

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(14)

    # Заголовки документа
    p_analitic_profile = doc.add_paragraph()
    p_analitic_profile.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_analitic_profile.paragraph_format.space_before = Pt(0)
    p_analitic_profile.paragraph_format.space_after = Pt(0)
    run_analitic_profile = p_analitic_profile.add_run("АНАЛІТИЧНИЙ ПРОФІЛЬ")
    run_analitic_profile.bold = True
    run_analitic_profile.font.size = Pt(14)

    p_on_person = doc.add_paragraph()
    p_on_person.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_on_person.paragraph_format.space_before = Pt(0)
    p_on_person.paragraph_format.space_after = Pt(0)
    run_on_person = p_on_person.add_run("на фізичну особу")
    run_on_person.bold = True
    run_on_person.font.size = Pt(14)

    doc.add_paragraph()

    if filled_blocks is None:
        filled_blocks = {}

    final_photo_bytes = None
    if dms_data and dms_data.get('photo_bytes'):
        final_photo_bytes = dms_data['photo_bytes']
    elif photo_bytes:
        final_photo_bytes = photo_bytes
    elif os.path.exists('default_avatar.png'):
        with open('default_avatar.png', 'rb') as f:
            final_photo_bytes = f.read()

    for block_name in EMPTY_DOSSIER_BLOCKS:
        if block_name == "АНКЕТНІ ДАНІ":
            if dms_data and dms_data.get('info'):
                append_dms_to_doc(doc, dms_data['info'], photo_bytes=final_photo_bytes, header_name="ІНФОРМАЦІЯ З ДМС")
            elif block_name in filled_blocks:
                content = filled_blocks[block_name]
                t_top = doc.add_table(rows=1, cols=1)
                t_top.width = Inches(6.5)
                cell_top = t_top.rows[0].cells[0]
                
                shd = OxmlElement('w:shd')
                shd.set(qn('w:fill'), '9BC2E6')
                cell_top._element.get_or_add_tcPr().append(shd)
                
                p_top = cell_top.paragraphs[0]
                p_top.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p_top.paragraph_format.space_before = Pt(0)
                p_top.paragraph_format.space_after = Pt(0)
                run_top = p_top.add_run("       АНКЕТНІ ДАНІ:")
                run_top.bold = True
                run_top.italic = True
                run_top.font.size = Pt(14)

                spacer = doc.add_paragraph()
                spacer.paragraph_format.space_before = Mm(3)
                spacer.paragraph_format.space_after = Mm(0)
                spacer.paragraph_format.line_spacing = 0

                table = doc.add_table(rows=1, cols=2)
                table.autofit = False
                
                left_cell = table.rows[0].cells[0]
                left_cell.width = Inches(2.0)
                if final_photo_bytes:
                    paragraph = left_cell.paragraphs[0]
                    run = paragraph.add_run()
                    run.add_picture(BytesIO(final_photo_bytes), width=Inches(1.8))

                right_cell = table.rows[0].cells[1]
                right_cell.width = Inches(4.5)
                right_cell.vertical_alignment = 1

                for para in right_cell.paragraphs:
                    para.clear()

                p = right_cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                run = p.add_run(content)
                run.font.size = Pt(14)
            else:
                t_top = doc.add_table(rows=1, cols=1)
                t_top.width = Inches(6.5)
                cell_top = t_top.rows[0].cells[0]
                
                shd = OxmlElement('w:shd')
                shd.set(qn('w:fill'), '9BC2E6')
                cell_top._element.get_or_add_tcPr().append(shd)
                
                p_top = cell_top.paragraphs[0]
                p_top.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p_top.paragraph_format.space_before = Pt(0)
                p_top.paragraph_format.space_after = Pt(0)
                run_top = p_top.add_run("       АНКЕТНІ ДАНІ:")
                run_top.bold = True
                run_top.italic = True
                run_top.font.size = Pt(14)

                spacer = doc.add_paragraph()
                spacer.paragraph_format.space_before = Mm(3)
                spacer.paragraph_format.space_after = Mm(0)
                spacer.paragraph_format.line_spacing = 0

                table = doc.add_table(rows=1, cols=2)
                table.autofit = False
                
                left_cell = table.rows[0].cells[0]
                left_cell.width = Inches(2.0)
                if final_photo_bytes:
                    paragraph = left_cell.paragraphs[0]
                    run = paragraph.add_run()
                    run.add_picture(BytesIO(final_photo_bytes), width=Inches(1.8))

                right_cell = table.rows[0].cells[1]
                right_cell.width = Inches(4.5)
                right_cell.vertical_alignment = 1

                for para in right_cell.paragraphs:
                    para.clear()

                for _ in range(5):
                    empty_p = right_cell.add_paragraph()
                    empty_p.paragraph_format.space_before = Pt(0)
                    empty_p.paragraph_format.space_after = Pt(0)

        elif block_name == "ДОКУМЕНТИ":
            if dms_data and dms_data.get('info') and dms_data['info'].get('documents'):
                add_block_header(doc, block_name)
                for doc_str in dms_data['info']['documents']:
                    p = doc.add_paragraph()
                    p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.space_after = Pt(2)
                    run = p.add_run(f"• {doc_str}")
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(14)
            elif block_name in filled_blocks:
                add_block_header(doc, block_name)
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(2)
                run = p.add_run(filled_blocks[block_name])
                run.font.name = 'Times New Roman'
                run.font.size = Pt(14)
            else:
                add_empty_block(doc, block_name)

        elif block_name == "НЕРУХОМЕ МАЙНО":
            if real_estate_data:
                append_real_estate_to_doc(doc, real_estate_data)
            elif block_name in filled_blocks:
                add_empty_block(doc, block_name)
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(2)
                run = p.add_run(filled_blocks[block_name])
                run.font.name = 'Times New Roman'
                run.font.size = Pt(14)
            else:
                add_empty_block(doc, block_name)

        elif block_name == "ТРАНСПОРТНІ ЗАСОБИ":
            if car_data:
                append_car_to_doc(doc, car_data)
            elif block_name in filled_blocks:
                add_empty_block(doc, block_name)
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(2)
                run = p.add_run(filled_blocks[block_name])
                run.font.name = 'Times New Roman'
                run.font.size = Pt(14)
            else:
                add_empty_block(doc, block_name)

        elif block_name == "РОДИННІ ЗВ'ЯЗКИ":
            if family_data:
                t = doc.add_table(rows=1, cols=1)
                t.width = Inches(6.5)
                cell = t.rows[0].cells[0]

                shading_elm = OxmlElement('w:shd')
                shading_elm.set(qn('w:fill'), '9BC2E6')
                cell._element.get_or_add_tcPr().append(shading_elm)

                tcPr = cell._element.get_or_add_tcPr()
                tcBorders = OxmlElement('w:tcBorders')
                for border in ['top', 'left', 'bottom', 'right']:
                    b = OxmlElement(f'w:{border}')
                    b.set(qn('w:val'), 'none')
                    tcBorders.append(b)
                tcPr.append(tcBorders)

                p_h = cell.paragraphs[0]
                p_h.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p_h.paragraph_format.space_before = Pt(0)
                p_h.paragraph_format.space_after = Pt(0)
                run_h = p_h.add_run("       РОДИННІ ЗВ'ЯЗКИ")
                run_h.bold = True
                run_h.italic = True
                run_h.font.size = Pt(14)

                for member in family_data:
                    relative_type = member.get('relative_type', 'Родич')

                    spacer = doc.add_paragraph()
                    spacer.paragraph_format.space_before = Mm(3)
                    spacer.paragraph_format.space_after = Mm(0)
                    spacer.paragraph_format.line_spacing = 0

                    if member.get('info'):
                        dms_info = member['info']
                        table = doc.add_table(rows=1, cols=2)
                        table.autofit = False

                        left_cell = table.rows[0].cells[0]
                        left_cell.width = Inches(1.8)
                        member_photo = member.get('photo_bytes')
                        if member_photo:
                            paragraph = left_cell.paragraphs[0]
                            run = paragraph.add_run()
                            run.add_picture(BytesIO(member_photo), width=Inches(1.6))
                        elif os.path.exists('default_avatar.png'):
                            with open('default_avatar.png', 'rb') as f:
                                paragraph = left_cell.paragraphs[0]
                                run = paragraph.add_run()
                                run.add_picture(BytesIO(f.read()), width=Inches(1.6))

                        right_cell = table.rows[0].cells[1]
                        right_cell.width = Inches(4.7)

                        p = right_cell.paragraphs[0]
                        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

                        r = p.add_run(f"[{relative_type}] ")
                        r.font.name = 'Times New Roman'
                        r.font.size = Pt(14)
                        r.bold = True

                        r = p.add_run(f"{dms_info.get('fio', '')}\n")
                        r.font.name = 'Times New Roman'
                        r.font.size = Pt(14)
                        r.bold = True

                        if dms_info.get('data'):
                            r = p.add_run(f"Дата народження: {dms_info['data']}\n")
                            r.font.name = 'Times New Roman'
                            r.font.size = Pt(14)

                        if dms_info.get('birthplace'):
                            r = p.add_run(f"Місце народження: {dms_info['birthplace']}\n")
                            r.font.name = 'Times New Roman'
                            r.font.size = Pt(14)

                        if dms_info.get('iphp'):
                            r = p.add_run(f"РНОКПП: {dms_info['iphp']}\n")
                            r.font.name = 'Times New Roman'
                            r.font.size = Pt(14)

                        if dms_info.get('uhzp'):
                            r = p.add_run(f"УНЗР: {dms_info['uhzp']}\n")
                            r.font.name = 'Times New Roman'
                            r.font.size = Pt(14)

                        if dms_info.get('tel'):
                            r = p.add_run(f"Телефон: {dms_info['tel']}")
                            r.font.name = 'Times New Roman'
                            r.font.size = Pt(14)
                            r.bold = True

                        # Адреса перебування - окремим параграфом під таблицею
                        if dms_info.get('adress'):
                            p_addr = doc.add_paragraph()
                            p_addr.paragraph_format.space_before = Pt(0)
                            p_addr.paragraph_format.space_after = Pt(2)
                            r_addr = p_addr.add_run(f"Адреса перебування: {dms_info['adress']}")
                            r_addr.font.name = 'Times New Roman'
                            r_addr.font.size = Pt(14)
                            r_addr.font.italic = True
                            r_addr.font.color.rgb = RGBColor(56, 86, 35)

                        # Паспортні дані - окремим списком під адресою
                        if dms_info.get('documents'):
                            p_docs = doc.add_paragraph()
                            p_docs.paragraph_format.space_before = Pt(0)
                            p_docs.paragraph_format.space_after = Pt(2)
                            r_docs = p_docs.add_run("ДОКУМЕНТИ:\n")
                            r_docs.bold = True
                            r_docs.font.name = 'Times New Roman'
                            r_docs.font.size = Pt(14)
                            for doc_str in dms_info['documents']:
                                r_doc = p_docs.add_run(f"• {doc_str}\n")
                                r_doc.font.name = 'Times New Roman'
                                r_doc.font.size = Pt(14)

                    elif member.get('manual_text'):
                        table = doc.add_table(rows=1, cols=2)
                        table.autofit = False

                        left_cell = table.rows[0].cells[0]
                        left_cell.width = Inches(1.8)
                        member_photo = member.get('photo_bytes')
                        if member_photo:
                            paragraph = left_cell.paragraphs[0]
                            run = paragraph.add_run()
                            run.add_picture(BytesIO(member_photo), width=Inches(1.6))
                        elif os.path.exists('default_avatar.png'):
                            with open('default_avatar.png', 'rb') as f:
                                paragraph = left_cell.paragraphs[0]
                                run = paragraph.add_run()
                                run.add_picture(BytesIO(f.read()), width=Inches(1.6))

                        right_cell = table.rows[0].cells[1]
                        right_cell.width = Inches(4.7)

                        p = right_cell.paragraphs[0]
                        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

                        r = p.add_run(f"[{relative_type}] ")
                        r.font.name = 'Times New Roman'
                        r.font.size = Pt(14)
                        r.bold = True

                        lines = member['manual_text'].split('\n')
                        for i, line in enumerate(lines):
                            if line.strip():
                                r = p.add_run(line.strip() + ("\n" if i < len(lines) - 1 else ""))
                                r.font.name = 'Times New Roman'
                                r.font.size = Pt(14)
            else:
                add_empty_block(doc, block_name)

        elif block_name == "ПЕРЕТИНИ ДЕРЖАВНОГО КОРДОНУ УКРАЇНИ":
            if border_crossing_data:
                append_border_crossing_to_doc(doc, border_crossing_data)
            else:
                add_empty_block(doc, block_name)

        elif block_name == "АДРЕСИ":
            if dms_data and dms_data.get('info') and dms_data['info'].get('adress'):
                add_block_header(doc, block_name)
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(2)
                run = p.add_run(f"Адреса перебування: {dms_data['info']['adress']}")
                run.font.name = 'Times New Roman'
                run.font.size = Pt(14)
            elif block_name in filled_blocks:
                add_block_header(doc, block_name)
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(2)
                run = p.add_run(filled_blocks[block_name])
                run.font.name = 'Times New Roman'
                run.font.size = Pt(14)
            else:
                add_empty_block(doc, block_name)

        elif block_name == "ПІДПРИЄМНИЦЬКА (ТРУДОВА) ДІЯЛЬНІСТЬ":
            # Пріоритет: 1. pension_data, 2. dms_data (ФОП), 3. filled_blocks, 4. порожній блок
            if pension_data and pension_data.get('formatted_line'):
                # Додаємо заголовок і текст одним блоком
                append_pension_to_doc(doc, pension_data, add_header=True)
            elif dms_data and dms_data.get('info') and dms_data['info'].get('fop'):
                add_block_header(doc, block_name)
                fop_data = dms_data['info']['fop']

                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(2)

                r = p.add_run("ІНФОРМАЦІЯ ПРО ФОП:\n")
                r.bold = True
                r.font.name = 'Times New Roman'
                r.font.size = Pt(14)

                r = p.add_run(f"ФОП: {fop_data['fio']}\n")
                r.font.name = 'Times New Roman'
                r.font.size = Pt(14)

                r = p.add_run(f"Статус: {fop_data['status']}\n")
                r.font.name = 'Times New Roman'
                r.font.size = Pt(14)

                r = p.add_run(f"Вид діяльності: {fop_data['kind_of_activity']}")
                r.font.name = 'Times New Roman'
                r.font.size = Pt(14)
            elif block_name in filled_blocks:
                add_block_header(doc, block_name)
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(2)
                run = p.add_run(filled_blocks[block_name])
                run.font.name = 'Times New Roman'
                run.font.size = Pt(14)
            else:
                add_empty_block(doc, block_name)

        else:
            if block_name in filled_blocks:
                add_block_header(doc, block_name)
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(2)
                run = p.add_run(filled_blocks[block_name])
                run.font.name = 'Times New Roman'
                run.font.size = Pt(14)
            else:
                add_empty_block(doc, block_name)

    section = doc.sections[0]

    header = section.header
    header.is_linked_to_previous = False

    header_para1 = header.paragraphs[0]
    header_para1.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    header_run1 = header_para1.add_run("УКА ГУНП")
    header_run1.font.name = "Times New Roman"
    header_run1.font.size = Pt(12)
    header_run1.font.color.rgb = RGBColor(255, 0, 0)

    header_para2 = header.add_paragraph()
    header_para2.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    today_date = datetime.today().strftime("%d.%m.%Y")
    header_run2 = header_para2.add_run(today_date)
    header_run2.font.name = "Times New Roman"
    header_run2.font.size = Pt(12)
    header_run2.font.color.rgb = RGBColor(255, 0, 0)

    footer = section.footer
    footer.is_linked_to_previous = False

    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    footer_run = footer_para.add_run("Управління кримінального аналізу")
    footer_run.font.name = "Times New Roman"
    footer_run.font.size = Pt(12)
    footer_run.font.bold = True

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()
