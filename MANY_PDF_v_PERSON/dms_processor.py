# -*- coding: utf-8 -*-
"""
Модуль для обробки PDF файлів Державної міграційної служби (ДМС)
"""

import fitz
import re
import requests
from bs4 import BeautifulSoup
from lxml import etree
from docx.shared import Inches, Pt, RGBColor, Mm, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import io
import os
import random
import datetime

def fop(ipn):
    """Перевірка статусу ФОП через YouControl"""
    if not ipn or ipn == 'невідомо':
        return False
        
    headers = {
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/107.0.0.0 Safari/537.36'
    }

    try:
        pages = requests.get(f'https://youcontrol.com.ua/search/?country=1&q={ipn}', headers=headers, timeout=10)
        if not pages.ok:
            return False

        soup = BeautifulSoup(pages.content, "html.parser")
        dom = etree.HTML(str(soup))
        
        fio = dom.xpath('//*[@id="catalog-company-file"]/div[2]/div[2]/div[2]/span')[0].text
        status = dom.xpath('//*[@id="catalog-company-file"]/div[2]/div[3]/div[2]/span/text()')[0].strip()
        kind_of_activity = dom.xpath('//*[@id="catalog-company-file"]/div[2]/div[6]/div[2]/div[2]/span')[0].text

        return {"fio": fio, "status": status, "kind_of_activity": kind_of_activity}
    except Exception:
        return False

def extract_dms_data(pdf_file):
    """
    Вилучає дані з PDF файлу ДМС.
    
    Args:
        pdf_file: UploadedFile object з Streamlit
        
    Returns:
        tuple: (dms_info_dict, photo_bytes, error_message)
    """
    try:
        # Відкриваємо PDF з байтів
        pdf_bytes = pdf_file.read()
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")

        srt_date = ''
        for page_num in range(doc.page_count):
            page = doc.load_page(page_num)
            srt_date += page.get_text()

        x = re.split("\n", srt_date)

        el = 'ІНФОРМАЦІЯ ПРО ОСОБУ'
        if el not in x:
            doc.close()
            return None, None, f"Файл {pdf_file.name} не містить очікувану інформацію ДМС"

        # Збір інформації про особу
        obj_people = {
            'fio': '', 'data': '', 'birthplace': '', 
            'tel': 'невідомо', 'adress': 'невідомо', 
            'uhzp': 'невідомо', 'iphp': 'невідомо',
            'documents': []
        }

        try:
            odj_inedx = x.index('Прізвище')
            obj_people['fio'] = f'{x[odj_inedx + 1]} {x[odj_inedx + 3]} {x[odj_inedx + 5]}'

            str_data = x[odj_inedx + 6].split(' ')
            obj_people['data'] = str_data[2] if len(str_data) > 2 else ''

            if 'Телефон' in x:
                odj_inedx = x.index('Телефон')
                obj_people['tel'] = x[odj_inedx + 1]

            if 'УНЗР' in x:
                odj_inedx = x.index('УНЗР')
                obj_people['uhzp'] = x[odj_inedx + 1]

            if 'РНОКПП' in x:
                odj_inedx = x.index('РНОКПП')
                obj_people['iphp'] = x[odj_inedx + 1]
        except (ValueError, IndexError):
            pass

        def get_address(vol, date_reper):
            try:
                index_start = vol.index(date_reper[0])
                index_end = vol.index(date_reper[1])
                addr = ''
                for n in range(index_start + 1, index_end - 1):
                    addr += vol[n] + ' '
                
                # Форматування адреси
                addr = addr.title()
                verification = ['М.', 'Вулиця', 'Район', 'Смт', 'Кв.', 'Буд.', 'Область', 'С.', 'Вул.', ' М ', "Пров.",
                              "Проспект.", "М-Н", "С-Ще", "Площа", "Просп."]
                
                for slovo in addr.split():
                    if re.search(r'\d{5}', slovo) is not None:
                        addr = addr.replace(slovo, '')

                for ver in verification:
                    addr = addr.replace(ver, ver.lower())

                return addr.replace('/', ', ').strip()
            except (ValueError, IndexError):
                return 'невідомо'

        obj_people['adress'] = get_address(x, ['перебування', 'Номер'])
        obj_people['birthplace'] = get_address(x, ['Місце народження', 'перебування'])

        # Обробка документів
        teloArr = ['Паспорт громадянина України', 
                  "Паспорт(и) громадянина України для виїзду за кордон",
                  'Свідоцтво про народження']

        def find_docs(arrX, index, doc_type_idx, exclude_indices):
            for w in range(index, len(arrX)):
                # Перевірка на початок іншого блоку документів
                is_other = False
                for ex_idx in exclude_indices:
                    if arrX[w] == teloArr[ex_idx]:
                        is_other = True
                        break
                if is_other: break
                
                if "Номер" == arrX[w]:
                    if w + 4 < len(arrX) and arrX[w + 3] == "Дійсний до:":
                        obj_people['documents'].append(f"{teloArr[doc_type_idx]} {arrX[w + 1]} дійсний до: {arrX[w + 4]}")
                    elif w + 5 < len(arrX) and arrX[w + 1] != "Дата видачі:":
                        obj_people['documents'].append(f"{teloArr[doc_type_idx]} {arrX[w + 1]} від {arrX[w + 3]} дійсний до: {arrX[w + 5]}")

        for i in range(len(teloArr)):
            if teloArr[i] in x:
                idx = x.index(teloArr[i])
                others = [j for j in range(len(teloArr)) if j != i]
                find_docs(x, idx, i, others)

        # Вилучення фото
        photo_bytes = None
        # Зазвичай фото на першій сторінці
        for img in doc.get_page_images(0):
            xref = img[0]
            pix = fitz.Pixmap(doc, xref)
            if pix.colorspace.n > 4: # Якщо CMYK, конвертуємо в RGB
                pix = fitz.Pixmap(fitz.csRGB, pix)
            
            photo_bytes = pix.tobytes("png")
            break

        doc.close()

        # Отримання інформації про ФОП
        obj_people['fop'] = fop(obj_people['iphp'])

        return obj_people, photo_bytes, None

    except Exception as e:
        return None, None, f"Помилка при обробці PDF ДМС: {str(e)}"

def append_dms_to_doc(doc, dms_info, photo_bytes=None, header_name="ІНФОРМАЦІЯ З ДМС"):
    """
    Додає блок інформації ДМС до документа.
    
    Args:
        doc: Document об'єкт
        dms_info: Словник з даними (результат extract_dms_data)
        photo_bytes: Байтовий рядок з фото
        header_name: Назва заголовка блоку
    """
    if not dms_info:
        return

    # Створюємо горизонтальну смугу-розділювач, якщо вказано заголовок
    if header_name:
        separator_table = doc.add_table(rows=1, cols=1)
        separator_table.width = Inches(6.5)
        separator_cell = separator_table.rows[0].cells[0]
        
        # Світло-блакитний фон (#9BC2E6)
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), '9BC2E6')
        separator_cell._element.get_or_add_tcPr().append(shading_elm)
        
        # Прибираємо границі
        tcPr = separator_cell._element.get_or_add_tcPr()
        tcBorders = OxmlElement('w:tcBorders')
        for border in ['top', 'left', 'bottom', 'right']:
            b = OxmlElement(f'w:{border}')
            b.set(qn('w:val'), 'none')
            tcBorders.append(b)
        tcPr.append(tcBorders)
        
        # Заголовок блоку
        p_separator = separator_cell.paragraphs[0]
        p_separator.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p_separator.paragraph_format.space_before = Pt(0)
        p_separator.paragraph_format.space_after = Pt(0)
        run_separator = p_separator.add_run("       " + header_name.upper())
        run_separator.bold = True
        run_separator.italic = True
        run_separator.font.size = Pt(14)
        run_separator.font.color.rgb = RGBColor(0, 0, 0)
        run_separator.font.name = 'Times New Roman'

    # Додаємо дані
    # Створюємо таблицю для фото та основної інформації
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Mm(3)
    spacer.paragraph_format.space_after = Mm(0)
    spacer.paragraph_format.line_spacing = 0
    
    info_table = doc.add_table(rows=1, cols=2)
    info_table.autofit = False
    
    left_cell = info_table.rows[0].cells[0]
    left_cell.width = Inches(1.8)
    
    right_cell = info_table.rows[0].cells[1]
    right_cell.width = Inches(4.7)
    
    # Фото в ліву клітинку
    if photo_bytes:
        p_img = left_cell.paragraphs[0]
        run_img = p_img.add_run()
        run_img.add_picture(io.BytesIO(photo_bytes), width=Inches(1.6))
    elif os.path.exists('default_avatar.png'):
        p_img = left_cell.paragraphs[0]
        run_img = p_img.add_run()
        run_img.add_picture('default_avatar.png', width=Inches(1.6))
        
    # Основна інформація в праву клітинку
    p = right_cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # ПІБ
    r = p.add_run(f"{dms_info['fio']}\n")
    r.font.name = 'Times New Roman'
    r.font.size = Pt(14)
    r.bold = True

    # Дата народження
    r = p.add_run(f"Дата народження: {dms_info['data']}\n")
    r.font.name = 'Times New Roman'
    r.font.size = Pt(14)

    # Місце народження
    r = p.add_run(f"Місце народження: {dms_info['birthplace']}\n")
    r.font.name = 'Times New Roman'
    r.font.size = Pt(14)

    # Додаткова інформація в ту ж клітинку таблиці
    r = p.add_run(f"РНОКПП: {dms_info['iphp']}\n")
    r.font.name = 'Times New Roman'
    r.font.size = Pt(14)

    r = p.add_run(f"УНЗР: {dms_info['uhzp']}\n")
    r.font.name = 'Times New Roman'
    r.font.size = Pt(14)

    r = p.add_run(f"Телефон: {dms_info['tel']}")
    r.font.name = 'Times New Roman'
    r.font.size = Pt(14)
    r.font.bold = True
