# -*- coding: utf-8 -*-
"""
Модуль для обробки даних про перетин кордону України з Excel файлів (ARKAN)
"""

from openpyxl import load_workbook
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime
import tempfile
import os


def process_excel_to_data(excel_file):
    """
    Обробка Excel файлу та витягування даних про перетин кордону.
    
    Args:
        excel_file: Завантажений файл з Streamlit (UploadedFile object)
        
    Returns:
        tuple: (exel_array, error_message)
            exel_array: список записів про перетин кордону
            error_message: текст помилки або None
    """
    
    # Зберігаємо Excel у тимчасову директорію
    with tempfile.NamedTemporaryFile(delete=False, suffix='.xlsx') as tmp_file:
        tmp_file.write(excel_file.getbuffer())
        excel_path = tmp_file.name
    
    try:
        wb = load_workbook(filename=excel_path)
        sheet_data = wb['Data']
        
        exel_array = []
        
        # Перевірка формату
        if sheet_data["AA3"].value == None:
            return None, f"Файл {excel_file.name} не містить очікувану структуру даних (AA3)"
        
        i = 3
        while sheet_data["A" + str(i)].value != None:
            if sheet_data["A" + str(i)].value == 'Скасовано':
                i += 1
                continue
            
            exel_str = []
            exel_str.append(sheet_data["A" + str(i)].value)  # 0 - Напрямок перетину
            exel_str.append(sheet_data["D" + str(i)].value)  # 1 - Громадянство
            exel_str.append(sheet_data["G" + str(i)].value)  # 2 - ПП перетину
            exel_str.append(sheet_data["I" + str(i)].value)  # 3 - Водій ТЗ
            exel_str.append(sheet_data["J" + str(i)].value.rstrip())  # 4 - Ділянка кордону
            exel_str.append(sheet_data["L" + str(i)].value)  # 5 - Дата, час перетину
            exel_str.append(sheet_data["M" + str(i)].value)  # 6 - ПІБ (українською)
            exel_str.append(sheet_data["N" + str(i)].value)  # 7 - ПІБ (латиницею)
            exel_str.append(sheet_data["P" + str(i)].value)  # 8 - Дата народження
            exel_str.append(sheet_data["S" + str(i)].value)  # 9 - Серія, номер документа
            exel_str.append(sheet_data["AB" + str(i)].value)  # 10 - Тип ПП
            exel_str.append(sheet_data["AE" + str(i)].value.rstrip())  # 11 - Вид ТЗ
            exel_str.append(sheet_data["AF" + str(i)].value.rstrip())  # 12 - Тип ТЗ
            exel_str.append(sheet_data["AH" + str(i)].value.rstrip())  # 13 - Марка ТЗ
            exel_str.append(sheet_data["AQ" + str(i)].value)  # 14 - Д/з номер
            exel_str.append(sheet_data["AR" + str(i)].value)  # 15 - VIN
            exel_str.append(sheet_data["H" + str(i)].value)  # 16 - Стать
            exel_str.append(sheet_data["T" + str(i)].value if sheet_data["T" + str(i)].value else "")  # 17 - Діти
            
            exel_array.append(exel_str)
            i += 1
        
        return exel_array, None
        
    except Exception as e:
        return None, f"Помилка обробки {excel_file.name}: {str(e)}"
    finally:
        # Видаляємо тимчасовий файл
        try:
            os.unlink(excel_path)
        except:
            pass


def append_border_crossing_to_doc(doc: Document, border_data: list):
    """
    Додає секцію з даними про перетин кордону до існуючого документа.
    
    Args:
        doc: Document об'єкт для додавання даних
        border_data: Масив даних про перетин кордону (exel_array)
    """
    
    if not border_data or len(border_data) == 0:
        return
    
    # Створюємо горизонтальну смугу на всю ширину сторінки
    separator_table = doc.add_table(rows=1, cols=1)
    separator_table.width = Inches(6.5)
    separator_cell = separator_table.rows[0].cells[0]
    
    # Налаштування світло-блакитного фону (#9BC2E6), як у інших блоків
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), '9BC2E6')
    separator_cell._element.get_or_add_tcPr().append(shading_elm)
    
    # Прибираємо границі
    tcPr = separator_cell._element.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for border in ['top', 'left', 'bottom', 'right']:
        b = OxmlElement(f'w:{border}')
        b.set(qn('w:val'), 'none')
        tcBorders.append(b)
    tcPr.append(tcBorders)
    
    # Додаємо текст на блакитній смузі (стиль: чорний, жирний, курсив, великі літери)
    p_separator = separator_cell.paragraphs[0]
    p_separator.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_separator.paragraph_format.space_before = Pt(0)
    p_separator.paragraph_format.space_after = Pt(0)
    run_separator = p_separator.add_run("       " + "ПЕРЕТИН КОРДОНУ УКРАЇНИ")
    run_separator.bold = True
    run_separator.italic = True
    run_separator.font.size = Pt(14)
    run_separator.font.color.rgb = RGBColor(0, 0, 0)  # Чорний текст
    run_separator.font.name = 'Times New Roman'

    # ПІБ та основна інформація
    fio_text = f"{border_data[0][6]} {border_data[0][8]}\n({border_data[0][7]})\n"
    fio = doc.add_paragraph()
    fio_text_1 = fio.add_run(fio_text)
    fio_text_1.bold = True
    fio_text_1.font.size = Pt(14)
    fio_text_1.font.name = 'Times New Roman'
    
    if border_data[0][11] == "Пішохід":
        text_TT = f"Пішохід п/п {border_data[0][2]} ділянка {border_data[0][4]}"
    else:
        text_TT = f"Заїхав п/п {border_data[0][2]} ділянка {border_data[0][4]} на {border_data[0][12]} {border_data[0][13]} {border_data[0][14]}"
    
    fio_text_2 = fio.add_run(f"Громадянин {border_data[0][1]}\nПАСПОРТ - {border_data[0][9]}")
    fio_text_2.bold = False
    fio_text_2.font.name = 'Times New Roman'
    fio_format = fio.paragraph_format
    fio_format.left_indent = Inches(3.5)
    
    # Статус перебування
    text_paragraf = doc.add_paragraph()
    if border_data[0][0] == "В`їзд":
        text = f"Знаходиться в Україні з {border_data[0][5]} \n ({text_TT})"
        color = 1
    else:
        text = f"Виїхав з України {border_data[0][5]} п/п {border_data[0][2]} ділянка {border_data[0][4]} на {border_data[0][12]} {border_data[0][13]} {border_data[0][14]}"
        color = 2
    
    if border_data[0][16] == "Чоловіча" and border_data[0][17] != "":
        text = f"""{text}
Підстава для виїзду - {border_data[0][17]}"""
    
    text_paragraf_3 = doc.add_paragraph()
    now = datetime.datetime.now()
    text_paragraf_3.add_run(f"Оперативна iнформацiя станом на {now.strftime('%d.%m.%Y %H:%M')} ")
    
    text_paragraf_1 = text_paragraf.add_run(text)
    text_paragraf_1.bold = True
    text_paragraf_1.font.size = Pt(14)
    text_paragraf_1.font.name = 'Times New Roman'
    
    if color == 1:
        text_paragraf_1.font.color.rgb = RGBColor(0, 100, 0)
    else:
        text_paragraf_1.font.color.rgb = RGBColor(178, 34, 34)
    
    text_paragraf_format = text_paragraf.paragraph_format
    text_paragraf_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    text_paragraf_format.space_before = Pt(18)
    text_paragraf_format.space_after = Pt(18)
    
    # Таблиця 1: Перетин кордону
    text = "1. Перетин кордону"
    text_paragraf_2 = doc.add_paragraph()
    text_paragraf_2_1 = text_paragraf_2.add_run(text)
    text_paragraf_2_1.bold = True
    text_paragraf_2_1.font.size = Pt(14)
    text_paragraf_2_1.font.name = 'Times New Roman'
    
    table_1 = doc.add_table(rows=1, cols=5)
    table_1_2 = table_1.rows[0].cells
    table_1.style = 'Light Grid'
    
    table_1_2[0].text = 'Дата'
    table_1_2[1].text = 'Напрямок'
    table_1_2[2].text = 'ПП перетину'
    table_1_2[3].text = 'Ділянка кордону'
    table_1_2[4].text = 'Тип ПП'
    
    for val in border_data:
        if val[11] == 'Автомобільний транспорт':
            val[11] = "aвто"
        if val[11] == 'Повітряний транспорт':
            val[11] = "лiтак"
        
        row_cells = table_1.add_row().cells
        row_cells[0].text = str(val[5])
        row_cells[1].text = str(val[0])
        row_cells[2].text = str(val[2])
        row_cells[3].text = str(val[4])
        row_cells[4].text = str(val[11])
    
    # Таблиця 2: Транспорт
    text_2 = "2. Tранспорт"
    text_paragraf_3 = doc.add_paragraph()
    text_paragraf_3_1 = text_paragraf_3.add_run(text_2)
    text_paragraf_3_1.bold = True
    text_paragraf_3_1.font.size = Pt(14)
    text_paragraf_3_1.font.name = 'Times New Roman'
    text_paragraf_format_2 = text_paragraf_3.paragraph_format
    text_paragraf_format_2.space_before = Pt(18)
    
    table_2 = doc.add_table(rows=1, cols=6)
    table_2_2 = table_2.rows[0].cells
    table_2.style = 'Light Grid'
    
    table_2_2[0].text = 'Дата'
    table_2_2[1].text = 'Напрямок'
    table_2_2[2].text = 'Водій'
    table_2_2[3].text = 'Транспорт'
    table_2_2[4].text = 'Модель'
    table_2_2[5].text = 'Д/з'
    table_2_2[0].width = Inches(1.3)
    
    for val in border_data:
        if val[11] == 'Пішохід':
            continue
        
        if val[12] == 'Легковий автомобіль':
            val[12] = 'Легковий'
        if val[12] == 'Літак пасажирський':
            val[12] = 'Літак'
        
        row_cells = table_2.add_row().cells
        row_cells[0].text = str(val[5])
        row_cells[1].text = str(val[0])
        row_cells[2].text = str(val[3])
        row_cells[3].text = str(val[12])
        row_cells[4].text = str(val[13])
        row_cells[5].text = str(val[14])
