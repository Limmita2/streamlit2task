import streamlit as st
import os
from openpyxl import load_workbook
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import datetime
import zipfile
from io import BytesIO
import tempfile

def process_excel(excel_file, temp_dir):
    """Обробка одного Excel файлу та створення Word документа"""

    # Зберігаємо Excel у тимчасову директорію
    excel_path = os.path.join(temp_dir, excel_file.name)
    with open(excel_path, 'wb') as f:
        f.write(excel_file.getbuffer())

    # Перевірка файлу
    if not (excel_path.endswith('.xlsx') or excel_path.endswith('.xls')):
        return None, f"Помилка: {excel_file.name} не є Excel файлом"

    try:
        wb = load_workbook(filename=excel_path)
        sheet_data = wb['Data']

        exel_array = []

        # Перевірка формату
        if sheet_data["AA3"].value == None:
            return None, f"Файл {excel_file.name} не містить очікувану структуру даних (AA3)"

        i = 3
        while sheet_data["A" + str(i)].value != None:
            if sheet_data["A" + str(i)].value == 'Скасовано':
                i += 1
                continue

            exel_str = []
            exel_str.append(sheet_data["A" + str(i)].value)  # 0 - Напрямок перетину
            exel_str.append(sheet_data["D" + str(i)].value)  # 1 - Громадянство
            exel_str.append(sheet_data["G" + str(i)].value)  # 2 - ПП перетину
            exel_str.append(sheet_data["I" + str(i)].value)  # 3 - Водій ТЗ
            exel_str.append(sheet_data["J" + str(i)].value.rstrip())  # 4 - Ділянка кордону
            exel_str.append(sheet_data["L" + str(i)].value)  # 5 - Дата, час перетину
            exel_str.append(sheet_data["M" + str(i)].value)  # 6 - ПІБ (українською)
            exel_str.append(sheet_data["N" + str(i)].value)  # 7 - ПІБ (латиницею)
            exel_str.append(sheet_data["P" + str(i)].value)  # 8 - Дата народження
            exel_str.append(sheet_data["S" + str(i)].value)  # 9 - Серія, номер документа
            exel_str.append(sheet_data["AB" + str(i)].value)  # 10 - Тип ПП
            exel_str.append(sheet_data["AE" + str(i)].value.rstrip())  # 11 - Вид ТЗ
            exel_str.append(sheet_data["AF" + str(i)].value.rstrip())  # 12 - Тип ТЗ
            exel_str.append(sheet_data["AH" + str(i)].value.rstrip())  # 13 - Марка ТЗ
            exel_str.append(sheet_data["AQ" + str(i)].value)  # 14 - Д/з номер
            exel_str.append(sheet_data["AR" + str(i)].value)  # 15 - VIN
            exel_str.append(sheet_data["H" + str(i)].value)  # 16 - Стать
            exel_str.append(sheet_data["T" + str(i)].value if sheet_data["T" + str(i)].value else "")  # 17 - Діти

            exel_array.append(exel_str)
            i += 1

        # Створення Word документа
        document = Document()
        sections = document.sections
        section = sections[0]
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(0.5)
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)

        # ПІБ та основна інформація
        fio_text = f"{exel_array[0][6]} {exel_array[0][8]}\n({exel_array[0][7]})\n"
        fio = document.add_paragraph()
        fio_text_1 = fio.add_run(fio_text)
        fio_text_1.bold = True
        fio_text_1.font.size = Pt(14)
        fio_text_1.font.name = 'Times New Roman'

        if exel_array[0][11] == "Пішохід":
            text_TT = f"Пішохід п/п {exel_array[0][2]} ділянка {exel_array[0][4]}"
        else:
            text_TT = f"Заїхав п/п {exel_array[0][2]} ділянка {exel_array[0][4]} на {exel_array[0][12]} {exel_array[0][13]} {exel_array[0][14]}"

        fio_text_2 = fio.add_run(f"Громадянин {exel_array[0][1]}\nПАСПОРТ - {exel_array[0][9]}")
        fio_text_2.bold = False
        fio_text_2.font.name = 'Times New Roman'
        fio_format = fio.paragraph_format
        fio_format.left_indent = Inches(3.5)

        # Статус перебування
        text_paragraf = document.add_paragraph()
        if exel_array[0][0] == "В`їзд":
            text = f"Знаходиться в Україні з {exel_array[0][5]} \n ({text_TT})"
            color = 1
        else:
            text = f"Виїхав з України {exel_array[0][5]} п/п {exel_array[0][2]} ділянка {exel_array[0][4]} на {exel_array[0][12]} {exel_array[0][13]} {exel_array[0][14]}"
            color = 2

        if exel_array[0][16] == "Чоловіча" and exel_array[0][17] != "":
            text = f"""{text}
Підстава для виїзду - {exel_array[0][17]}"""

        text_paragraf_3 = document.add_paragraph()
        now = datetime.datetime.now()
        text_paragraf_3.add_run(f"Оперативна iнформацiя станом на {now.strftime('%d.%m.%Y %H:%M')} ")

        text_paragraf_1 = text_paragraf.add_run(text)
        text_paragraf_1.bold = True
        text_paragraf_1.font.size = Pt(14)
        text_paragraf_1.font.name = 'Times New Roman'

        if color == 1:
            text_paragraf_1.font.color.rgb = RGBColor(0, 100, 0)
        else:
            text_paragraf_1.font.color.rgb = RGBColor(178, 34, 34)

        text_paragraf_format = text_paragraf.paragraph_format
        text_paragraf_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        text_paragraf_format.space_before = Pt(18)
        text_paragraf_format.space_after = Pt(18)

        # Таблиця 1: Перетин кордону
        text = "1. Перетин кордону"
        text_paragraf_2 = document.add_paragraph()
        text_paragraf_2_1 = text_paragraf_2.add_run(text)
        text_paragraf_2_1.bold = True
        text_paragraf_2_1.font.size = Pt(14)
        text_paragraf_2_1.font.name = 'Times New Roman'

        table_1 = document.add_table(rows=1, cols=5)
        table_1_2 = table_1.rows[0].cells
        table_1.style = 'Light Grid'

        table_1_2[0].text = 'Дата'
        table_1_2[1].text = 'Напрямок'
        table_1_2[2].text = 'ПП перетину'
        table_1_2[3].text = 'Ділянка кордону'
        table_1_2[4].text = 'Тип ПП'

        for val in exel_array:
            if val[11] == 'Автомобільний транспорт':
                val[11] = "aвто"
            if val[11] == 'Повітряний транспорт':
                val[11] = "лiтак"

            row_cells = table_1.add_row().cells
            row_cells[0].text = str(val[5])
            row_cells[1].text = str(val[0])
            row_cells[2].text = str(val[2])
            row_cells[3].text = str(val[4])
            row_cells[4].text = str(val[11])

        # Таблиця 2: Транспорт
        text_2 = "2. Tранспорт"
        text_paragraf_3 = document.add_paragraph()
        text_paragraf_3_1 = text_paragraf_3.add_run(text_2)
        text_paragraf_3_1.bold = True
        text_paragraf_3_1.font.size = Pt(14)
        text_paragraf_3_1.font.name = 'Times New Roman'
        text_paragraf_format_2 = text_paragraf_3.paragraph_format
        text_paragraf_format_2.space_before = Pt(18)

        table_2 = document.add_table(rows=1, cols=6)
        table_2_2 = table_2.rows[0].cells
        table_2.style = 'Light Grid'

        table_2_2[0].text = 'Дата'
        table_2_2[1].text = 'Напрямок'
        table_2_2[2].text = 'Водій'
        table_2_2[3].text = 'Транспорт'
        table_2_2[4].text = 'Модель'
        table_2_2[5].text = 'Д/з'
        table_2_2[0].width = Inches(1.3)

        for val in exel_array:
            if val[11] == 'Пішохід':
                continue

            if val[12] == 'Легковий автомобіль':
                val[12] = 'Легковий'
            if val[12] == 'Літак пасажирський':
                val[12] = 'Літак'

            row_cells = table_2.add_row().cells
            row_cells[0].text = str(val[5])
            row_cells[1].text = str(val[0])
            row_cells[2].text = str(val[3])
            row_cells[3].text = str(val[12])
            row_cells[4].text = str(val[13])
            row_cells[5].text = str(val[14])

        # Збереження Word документа
        docx_filename = f"{exel_array[0][7]}.docx"
        docx_path = os.path.join(temp_dir, docx_filename)
        document.save(docx_path)

        return docx_path, None

    except Exception as e:
        return None, f"Помилка обробки {excel_file.name}: {str(e)}"

# Streamlit інтерфейс
st.set_page_config(page_title="Excel to Word Converter", page_icon="📊", layout="centered")

st.title("📊 Конвертер Excel → Word")
st.write("Завантажте один або декілька Excel файлів для створення Word документів")


# Завантаження файлів
uploaded_files = st.file_uploader(
    "Перетягніть файли сюди або натисніть для вибору",
    type=['xlsx', 'xls'],
    accept_multiple_files=True,
    help="Можна завантажити декілька Excel файлів одночасно"
)

if uploaded_files:
    st.info(f"Завантажено файлів: {len(uploaded_files)}")

    # Кнопка обробки
    if st.button("🔄 Обробити", type="primary", use_container_width=True):
        with st.spinner("Обробка файлів..."):
            # Створюємо тимчасову директорію
            with tempfile.TemporaryDirectory() as temp_dir:
                processed_files = []
                errors = []

                # Прогрес бар
                progress_bar = st.progress(0)
                status_text = st.empty()

                for idx, uploaded_file in enumerate(uploaded_files):
                    status_text.text(f"Обробка: {uploaded_file.name}")

                    docx_path, error = process_excel(uploaded_file, temp_dir)

                    if error:
                        errors.append(error)
                    elif docx_path:
                        processed_files.append(docx_path)

                    progress_bar.progress((idx + 1) / len(uploaded_files))

                status_text.empty()
                progress_bar.empty()

                # Показуємо помилки
                if errors:
                    st.error("Помилки при обробці:")
                    for error in errors:
                        st.write(f"❌ {error}")

                # Завантаження результатів
                if processed_files:
                    st.success(f"✅ Успішно оброблено: {len(processed_files)} файл(ів)")

                    if len(processed_files) == 1:
                        # Один файл - завантажуємо напряму
                        with open(processed_files[0], 'rb') as f:
                            docx_data = f.read()

                        filename = os.path.basename(processed_files[0])
                        st.download_button(
                            label="💾 Завантажити Word документ",
                            data=docx_data,
                            file_name=filename,
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            use_container_width=True
                        )
                    else:
                        # Декілька файлів - створюємо архів
                        zip_buffer = BytesIO()
                        with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
                            for docx_path in processed_files:
                                zip_file.write(docx_path, os.path.basename(docx_path))

                        st.download_button(
                            label=f"💾 Завантажити всі файли ({len(processed_files)} шт.)",
                            data=zip_buffer.getvalue(),
                            file_name="border_crossing_documents.zip",
                            mime="application/zip",
                            use_container_width=True
                        )
                else:
                    st.warning("Не вдалося обробити жоден файл")

st.markdown("---")
st.caption("Конвертер Excel → Word | Формування звітів про перетин кордону")
