import streamlit as st
import os
import fitz
import re
import random
import requests
from bs4 import BeautifulSoup
from lxml import etree
from docx.shared import Inches, Pt, RGBColor, Cm
import docx
import zipfile
from io import BytesIO
import tempfile

def fop(ipn):
    headers = {
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/107.0.0.0 Safari/537.36'
    }

    pages = requests.get(f'https://youcontrol.com.ua/search/?country=1&q={ipn}', headers=headers)
    if not pages.ok:
        return False

    soup = BeautifulSoup(pages.content, "html.parser")
    dom = etree.HTML(str(soup))
    try:
        fio = dom.xpath('//*[@id="catalog-company-file"]/div[2]/div[2]/div[2]/span')[0].text
        status = dom.xpath('//*[@id="catalog-company-file"]/div[2]/div[3]/div[2]/span/text()')[0].strip()
        kind_of_activity = dom.xpath('//*[@id="catalog-company-file"]/div[2]/div[6]/div[2]/div[2]/span')[0].text
    except Exception:
        return False

    youcontrol = {"fio": fio, "status": status, "kind_of_activity": kind_of_activity}
    return youcontrol

def process_pdf(pdf_file, temp_dir):
    """Обробка одного PDF файлу"""

    # Зберігаємо PDF у тимчасову директорію
    pdf_path = os.path.join(temp_dir, pdf_file.name)
    with open(pdf_path, 'wb') as f:
        f.write(pdf_file.getbuffer())

    # Перевірка файлу
    if not pdf_path.endswith('.pdf'):
        return None, f"Помилка: {pdf_file.name} не є PDF файлом"

    try:
        doc = fitz.open(pdf_path)

        srt_date = ''
        for page_num in range(doc.page_count):
            page = doc.load_page(page_num)
            srt_date += page.get_text()

        x = re.split("\n", srt_date)

        el = 'ІНФОРМАЦІЯ ПРО ОСОБУ'
        if el not in x:
            doc.close()
            return None, f"Файл {pdf_file.name} не містить очікувану інформацію"

        # Збір інформації про особу
        obj_people = {
            'fio': '', 'data': '', 'birthplace': '', 
            'tel': 'невідомо', 'adress': 'невідомо', 
            'uhzp': 'невідомо', 'iphp': 'невідомо'
        }

        odj_inedx = x.index('Прізвище')
        obj_people['fio'] = f'{x[odj_inedx + 1]} {x[odj_inedx + 3]} {x[odj_inedx + 5]}'

        str_data = x[odj_inedx + 6].split(' ')
        obj_people['data'] = str_data[2]

        if 'Телефон' in x:
            odj_inedx = x.index('Телефон')
            obj_people['tel'] = x[odj_inedx + 1]

        odj_inedx = x.index('УНЗР')
        obj_people['uhzp'] = x[odj_inedx + 1]

        odj_inedx = x.index('РНОКПП')
        obj_people['iphp'] = x[odj_inedx + 1]

        def adress(vol, date_reper):
            index_reper = [vol.index(date_reper[0]), vol.index(date_reper[1])]
            adres = ''
            for n in range(index_reper[0] + 1, index_reper[1] - 1):
                adres += vol[n] + ' '
            adres = adres.title()

            verification = ['М.', 'Вулиця', 'Район', 'Смт', 'Кв.', 'Буд.', 'Область', 'С.', 'Вул.', ' М ', "Пров.",
                          "Проспект.", "М-Н", "С-Ще", "Площа", "Просп."]

            for slovo in adres.split():
                if re.search('\d{5}', slovo) is not None:
                    adres = adres.replace(slovo, '')

            for ver in verification:
                adres = adres.replace(ver, ver.lower())

            adres = adres.replace('/', ', ')
            return adres.strip()

        obj_people['adress'] = adress(x, ['перебування', 'Номер'])
        obj_people['birthplace'] = adress(x, ['Місце народження', 'перебування'])

        # Обробка документів
        text = []
        teloArr = ['Паспорт громадянина України', 
                  "Паспорт(и) громадянина України для виїзду за кордон",
                  'Свідоцтво про народження']

        def povtorPoisk(arrX, index, telo, minustelo):
            for w in range(index, len(arrX)):
                if arrX[w] == teloArr[minustelo[0]]:
                    break
                elif arrX[w] == teloArr[minustelo[1]]:
                    break
                elif "Номер" == arrX[w] and arrX[w + 3] == "Дійсний до:":
                    text.append(f"{teloArr[telo]} {arrX[w + 1]} дійсний до: {arrX[w + 4]}")
                elif "Номер" == arrX[w] and arrX[w + 1] != "Дата видачі:":
                    text.append(f"{teloArr[telo]} {arrX[w + 1]} від {arrX[w + 3]} дійсний до: {arrX[w + 5]}")

        def pgu(x_arr, telo, minustelo):
            if teloArr[telo] not in x_arr:
                return
            index = x_arr.index(teloArr[telo])
            povtorPoisk(x_arr, index, telo, minustelo)

        pgu(x, 0, [1, 2])
        pgu(x, 1, [0, 2])
        pgu(x, 2, [0, 1])

        # Збереження фото
        path_foto = None
        for i in range(1):
            for img in doc.get_page_images(i):
                xref = img[0]
                pix = fitz.Pixmap(doc, xref)
                pix1 = fitz.Pixmap(fitz.csRGB, pix)
                number = random.randint(0, 1000)
                path_foto = os.path.join(temp_dir, f'image_{number}.png')
                pix1.save(path_foto)
                break

        doc.close()

        # Отримання інформації про ФОП
        fop_fio = fop(obj_people['iphp'])

        # Створення Word документа
        docx_doc = docx.Document()
        sections = docx_doc.sections
        section = sections[0]
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(0.5)
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)

        style = docx_doc.styles['Normal']
        style.font.name = 'Times New Roman'
        style.font.size = Pt(14)

        run = docx_doc.add_paragraph().add_run(obj_people['fio'].title())
        run.font.color.rgb = RGBColor(0, 32, 96)
        run.font.bold = True

        if path_foto:
            docx_doc.add_picture(path_foto, width=Cm(3))

        paragr = docx_doc.add_paragraph()
        paragr.add_run(f"{obj_people['data']} р.н.")
        paragr.add_run(', місце народження: ')
        paragr.add_run(f"{obj_people['birthplace']}\n")
        paragr.add_run(f"РНОКПП: ")
        paragr.add_run(f"{obj_people['iphp']}\n")
        paragr.add_run(f"{'\n'.join(text)}")
        paragr.add_run(f"\nУНЗР: ")
        paragr.add_run(f"{obj_people['uhzp']}\n")
        paragr.add_run(f"Можливе місце проживання: ")
        ruta = paragr.add_run(f"{obj_people['adress']}\n")
        ruta.font.color.rgb = RGBColor(56, 86, 35)
        ruta.font.italic = True
        paragr.add_run(f"Користується абонентським номером: ")
        paragr.add_run(f"{obj_people['tel']}\n").bold = True

        if fop_fio:
            fop_pag = docx_doc.add_paragraph()
            fop_pag.add_run(f'ФОП ')
            fop_pag.add_run(f"{fop_fio['fio']}").bold = True
            fop_pag.add_run(f", статус: {fop_fio['status']}, Основний вид діяльноcті: {fop_fio['kind_of_activity']}.")

        # Збереження DOCX
        docx_filename = f"{obj_people['fio']}.docx"
        docx_path = os.path.join(temp_dir, docx_filename)
        docx_doc.save(docx_path)

        return docx_path, None

    except Exception as e:
        return None, f"Помилка обробки {pdf_file.name}: {str(e)}"

# Streamlit інтерфейс
st.set_page_config(page_title="PDF to DOCX Converter", page_icon="📄", layout="centered")

st.title("📄 Конвертер PDF → DOCX")
st.write("Завантажте один або декілька PDF файлів для конвертації")

# Завантаження файлів
uploaded_files = st.file_uploader(
    "Перетягніть файли сюди або натисніть для вибору",
    type=['pdf'],
    accept_multiple_files=True,
    help="Можна завантажити декілька PDF файлів одночасно"
)

if uploaded_files:
    st.info(f"Завантажено файлів: {len(uploaded_files)}")

    # Кнопка обробки
    if st.button("🔄 Обробити", type="primary", use_container_width=True):
        with st.spinner("Обробка файлів..."):
            # Створюємо тимчасову директорію
            with tempfile.TemporaryDirectory() as temp_dir:
                processed_files = []
                errors = []

                # Прогрес бар
                progress_bar = st.progress(0)
                status_text = st.empty()

                for idx, uploaded_file in enumerate(uploaded_files):
                    status_text.text(f"Обробка: {uploaded_file.name}")

                    docx_path, error = process_pdf(uploaded_file, temp_dir)

                    if error:
                        errors.append(error)
                    elif docx_path:
                        processed_files.append(docx_path)

                    progress_bar.progress((idx + 1) / len(uploaded_files))

                status_text.empty()
                progress_bar.empty()

                # Показуємо помилки
                if errors:
                    st.error("Помилки при обробці:")
                    for error in errors:
                        st.write(f"❌ {error}")

                # Завантаження результатів
                if processed_files:
                    st.success(f"✅ Успішно оброблено: {len(processed_files)} файл(ів)")

                    if len(processed_files) == 1:
                        # Один файл - завантажуємо напряму
                        with open(processed_files[0], 'rb') as f:
                            docx_data = f.read()

                        filename = os.path.basename(processed_files[0])
                        st.download_button(
                            label="💾 Завантажити DOCX",
                            data=docx_data,
                            file_name=filename,
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            use_container_width=True
                        )
                    else:
                        # Декілька файлів - створюємо архів
                        zip_buffer = BytesIO()
                        with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
                            for docx_path in processed_files:
                                zip_file.write(docx_path, os.path.basename(docx_path))

                        st.download_button(
                            label=f"💾 Завантажити всі файли ({len(processed_files)} шт.)",
                            data=zip_buffer.getvalue(),
                            file_name="converted_documents.zip",
                            mime="application/zip",
                            use_container_width=True
                        )
                else:
                    st.warning("Не вдалося обробити жоден файл")

st.markdown("---")
st.caption("Конвертер PDF → DOCX | Зберігає форматування та зображення")
